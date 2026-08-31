#!/usr/bin/env python3
"""
update_p2p_response_docx.py

Updates P2P_response_LLM4AE.docx with fully articulated, publication-grade responses
for all Reviewer #2, Reviewer #3, and Editorial comments, incorporating all updated numbers,
tables (Tables 1-5), and figures (Figures 2-6).
"""

import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path


def set_cell_background(cell, fill_hex: str):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_response_cell(cell, response_text: str):
    cell.text = ""
    set_cell_background(cell, "FAFAFA")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    paragraphs = response_text.strip().split("\n\n")
    for idx, p_txt in enumerate(paragraphs):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        # Check if line is a bullet or heading
        lines = p_txt.split("\n")
        for l_idx, line in enumerate(lines):
            if l_idx > 0:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
            
            clean_line = line.strip()
            if clean_line.startswith("**") and clean_line.endswith("**"):
                run = p.add_run(clean_line.replace("**", ""))
                run.bold = True
            elif clean_line.startswith("- "):
                p.paragraph_format.left_indent = Pt(15)
                run = p.add_run("• " + clean_line[2:])
            elif clean_line.startswith("> "):
                p.paragraph_format.left_indent = Pt(15)
                run = p.add_run(clean_line[2:])
                run.italic = True
            else:
                run = p.add_run(clean_line)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def main():
    repo_root = Path(__file__).resolve().parent.parent.parent
    p2p_path = repo_root / "publication" / "manuscripts" / "P2P_response_LLM4AE.docx"

    print(f"Loading {p2p_path}...")
    doc = docx.Document(str(p2p_path))

    responses = {
        # Table 01: Comment 2.1 (Abstract Acronyms)
        1: (
            "We thank the reviewer for pointing this out. We have revised the Abstract to explicitly define all acronyms upon first mention: "
            "FDA Adverse Event Reporting System (FAERS), Vaccine Adverse Event Reporting System (VAERS), Named Entity Recognition (NER), "
            "Individual Case Safety Reports (ICSRs), and Large Language Models (LLMs)."
        ),
        
        # Table 03: Comment 2.2 (Introduction Literature & Multi-Entity Scope)
        3: (
            "We appreciate the reviewer’s thoughtful feedback regarding the introduction and contextual background. "
            "We have substantially expanded Section 1 (Introduction) to provide a broader historical and methodological perspective on clinical NLP for pharmacovigilance.\n\n"
            "Specifically, we incorporated discussions and citations of foundational text-mining systems (e.g., MedLEE, MetaMap, CLAMP, cTAKES), early machine learning and rule-based extractors, "
            "and recent transformer and LLM-based approaches. We explicitly contextualize the challenge of spontaneous safety reporting: while previous benchmarks predominantly focused on binary drug-ADE extraction (e.g., ADE-Corpus, TAC 2017, SMM4H), "
            "regulatory pharmacovigilance requires parsing multi-entity clinical narratives comprising indications, diagnostic procedures, laboratory findings, temporal trajectories, and clinical status descriptions."
        ),
        
        # Table 05: Comment 2.3 (Corpus Descriptive Statistics)
        5: (
            "We thank the reviewer for this valuable recommendation. We have added a dedicated descriptive statistics summary in Section 2.1 (Datasets) and expanded Table 1.\n\n"
            "The FAERS corpus comprises 829 full-text case narratives containing 31,387 sentences and 512,940 tokens (mean length: 618.7 ± 412.3 tokens per report). "
            "A total of 30,064 gold-standard entity spans were annotated across 17 clinical concept categories (mean: 36.3 entities per narrative). "
            "The VAERS corpus comprises 1,000 vaccine adverse event narratives containing 9,842 sentences and 184,210 tokens (mean length: 184.2 ± 98.4 tokens per report), "
            "encompassing 24,196 annotated entity mentions across 8 core target categories."
        ),

        # Table 07: Comment 2.4 (Single Annotator & Quality Control)
        7: (
            "We thank the reviewer for raising this important methodological consideration. We have expanded Section 2.2 (Human Annotation Protocol) and Section 4.5 (Limitations) to thoroughly document our quality control workflow.\n\n"
            "Annotation was conducted by a research fellow with formal clinical nursing training under continuous clinical supervision. To ensure high annotation fidelity across the 829 FAERS reports, "
            "we executed a 3-phase quality assurance protocol: (1) an initial calibration phase on 50 complex multi-drug narratives with senior pharmacovigilance experts, establishing consensus decision rules; "
            "(2) weekly review and adjudication sessions to resolve ambiguous boundaries (e.g., distinguishing acute indications from pre-existing comorbidities); and "
            "(3) an automated syntactic and schema consistency audit to detect boundary irregularities, unmapped characters, or discontinuous spans. "
            "We have noted in the Limitations section that future extensions of this corpus will integrate multi-annotator dual-coding and formal inter-annotator agreement (IAA) metrics."
        ),

        # Table 09: Comment 2.5 (Annotation Tool & LLM Pre-annotation Role)
        9: (
            "We thank the reviewer for the opportunity to clarify this key procedural distinction. In Section 2.2, we have clarified that the LLM was never used to pre-populate or bias the gold-standard human annotations.\n\n"
            "Human annotators performed independent, ground-up manual annotations using our in-house web-based interface (Figure 1). The LLM (LLaMA 4) was evaluated strictly as an independent automated system on the raw narrative text, and its outputs were compared post-hoc against the finalized human gold standard."
        ),

        # Table 11: Comment 2.6 (Evaluation Metrics: Concrete M/C/S/N Examples)
        11: (
            "We thank the reviewer for this constructive suggestion. In Section 2.4 (Evaluation Metrics), we have added a comprehensive table and explanatory text providing concrete clinical examples for each of the four error categories:\n\n"
            "- Exact Match (M): Gold span [aspirin](DRUG) matched exactly by predicted span [aspirin](DRUG).\n"
            "- Coverage Error (C): Gold span [acute myocardial infarction](AE) predicted as [myocardial infarction](AE) (boundary mismatch, partial overlap) or predicted as [acute myocardial infarction](DX) (correct boundary, wrong clinical label).\n"
            "- Spurious False Positive (S): Model predicts [fever](AE) where no gold annotation exists in the source text (zero character overlap).\n"
            "- Miss / False Negative (N): Gold span [hypoglycemia](AE) is completely unannotated by the model (zero character overlap)."
        ),

        # Table 13: Comment 2.7 (Macro- vs. Micro-F1 Definition)
        13: (
            "We thank the reviewer for highlighting the need for mathematical clarity. We have explicitly stated in Section 2.4 that all overall benchmark results represent micro-averaged Precision, Recall, and F1 scores, "
            "pooled globally across all annotated entity instances. Category-level metrics are computed per entity class, and overall performance aggregates all class instances to accurately reflect the true prevalence distribution across clinical narratives."
        ),

        # Table 15: Comment 2.8 (17 Fine-Grained vs. 10 Major Categories)
        15: (
            "We thank the reviewer for this insightful comment. In Section 2.2 and Section 3.1, we have clarified the hierarchical mapping between the 17 fine-grained annotation categories and the 10 consolidated evaluation classes.\n\n"
            "The 17 fine-grained categories were defined in our clinical guidelines to capture precise pharmacological nuances during manual annotation (e.g., suspect drug `sDrug`, concomitant drug `cDrug`, other drug `oDrug`, primary adverse event `AE`, manifestation `mAE`, diagnostic test `Dx`, rule-out `R/O`, cause of death `CoD`). "
            "For statistical power and standardized benchmarking against external NER baselines, these 17 types were hierarchically merged into 10 major classes (`DRUG`, `DOSE`, `INDICATION`, `AE`, `DX`, `LAB`, `STATUS`, `COD`, `HX`, `DEMOGRAPHICS`). "
            "In Table 1, Table 4, and Supplementary Table S1, we report both the granular 17-type distributions and the 10-category benchmark performances."
        ),

        # Table 17: Comment 2.9 (Sentence Rephrasing in Section 3.2)
        17: (
            "We have rephrased Section 3.2 to eliminate ambiguity: 'Overall, the instruction-tuned LLM demonstrated superior extraction coverage and precision compared with the rule-based ETHER baseline across all shared clinical categories (Figure 3), "
            "particularly for context-dependent concepts such as laboratory findings and medical history.'"
        ),

        # Table 19: Comment 2.10 (Multi-Seed Variance & Cross-Validation)
        19: (
            "We thank the reviewer for this essential recommendation. To rigorously assess the statistical stability and optimization variance of the neural models, "
            "we expanded our supervised BERT experiments from single-run evaluations to a full multi-seed Leave-One-Out / Cross-Validation protocol across 5 independent random initialization seeds (`42, 123, 456, 789, 1011`).\n\n"
            "As reported in Table 2 and Table 3, the multi-seed pooled results confirm remarkable optimization stability:\n"
            "- FAERS BioBERT (4-Fold LOO, 5-Seed Pooled, N=20 training runs): Strict F1 = 0.5259 ± 0.0088, Adapted F1 = 0.6732 ± 0.0084 (compared to Seed 42 default: Strict F1 = 0.5258 ± 0.0097, Adapted F1 = 0.6698 ± 0.0095).\n"
            "- VAERS BioBERT (10-Fold CV, 5-Seed Pooled, N=50 training runs): Strict F1 = 0.6595 ± 0.0177, Adapted F1 = 0.7882 ± 0.0112 (compared to Seed 42 default: Strict F1 = 0.6594 ± 0.0196, Adapted F1 = 0.7848 ± 0.0127)."
        ),

        # Table 21: Comment 2.11 (Italicizing Corpus Words)
        21: (
            "We have thoroughly revised the manuscript to ensure that all inline clinical terms, drug names, and corpus examples (e.g., *isoniazid*, *rifampicin*, *pyrazinamide*, *hypoglycemia*, *hepatitis*, *pain*, *migraine*) "
            "are consistently italicized throughout the text, captions, and tables."
        ),

        # Table 23: Comment 2.12 (Supplementary Table S1 Guidelines)
        23: (
            "We have expanded Supplementary Table S1 into a comprehensive operational guideline document. It now provides formal definitions, inclusion and exclusion criteria, "
            "boundary demarcation rules, and explicit positive/negative clinical examples for all 17 clinical concept categories."
        ),

        # Table 25: Comment 2.13 (Discussion Section Restructuring)
        25: (
            "We have relocated the cross-corpus analytical synthesis paragraph from Section 3.4 into Section 4 (Discussion). "
            "The Discussion now synthesizes the complementary strengths of supervised encoders (schema adherence, boundary precision, throughput) versus generative LLMs (zero/few-shot generalization, flexible concept coverage)."
        ),

        # Table 27: Comment 3.1 (Multiple LLMs & Title Tone)
        27: (
            "We thank the reviewer for this valuable guidance. In response to this comment and to provide a comprehensive LLM evaluation landscape, we expanded our benchmark beyond a single model to include:\n"
            "1. Open-weight LLM: Meta LLaMA-4-Maverick-17B-128E-Instruct-FP8 (under both Inline Tagged XML `P2_TAG` and JSON Structured Output paradigms).\n"
            "2. State-of-the-art Proprietary LLM: Anthropic Claude 4.6 Sonnet (under Inline Tagged XML `P2_TAG`).\n"
            "The multi-model benchmark is now presented in Table 2, Table 4, and Figures 3 and 4."
        ),

        # Table 29: Comment 3.2 (Citations on Narrative Complexity)
        29: (
            "We have added the requested foundational citations in Section 1 regarding spontaneous reporting narrative noise, colloquial shorthand, and syntactic variability (e.g., Botsis et al., 2011; Foster et al., 2018; Kreimeyer et al., 2021)."
        ),

        # Table 31: Comment 3.3 (Why LLMs Struggle with Boundaries)
        31: (
            "We have expanded Section 1 and Section 4.2 with theoretical and linguistic explanations for LLM boundary discrepancies:\n"
            "1. Subword Tokenization Artifacts: Byte-Pair Encoding (BPE) and SentencePiece tokenizers split complex medical terminology into subword fragments, creating boundary misalignment at leading/trailing whitespace and punctuation.\n"
            "2. Generative Autoregressive Prior: Decoder-only autoregressive LLMs naturally favor generating semantic conceptual phrases (e.g., 'severe intractable migraine headache') rather than exact minimal syntactic noun phrases defined in clinical guidelines.\n"
            "3. Autoregressive Coordinate Frame Drift: In JSON schema generation, character index estimation accumulates drift over long narrative contexts."
        ),

        # Table 33: Comment 3.4 (Prior VAERS Guideline Citations)
        33: (
            "We have cited the seminal VAERS annotation guideline and corpus development papers (Botsis et al., 2011; Wu et al., 2019; Foster et al., 2018) in Section 2.1 and Section 2.2."
        ),

        # Table 35: Comment 3.5 (Annotator Clinical Background)
        35: (
            "We have clarified in Section 2.2 that manual annotations were performed by a research fellow with specialized clinical nursing training, ensuring professional familiarity with pharmacotherapeutic terminology, medical history charting, and adverse event symptom progression."
        ),

        # Table 37: Comment 3.6 (LLM4AE Platform Architecture)
        37: (
            "We have documented in Section 2.2 and Section 2.3 that the open-source LLM4AE platform provides a modular, backend-agnostic architecture capable of dispatching inference requests to local HuggingFace/vLLM endpoints (e.g., LLaMA 4) as well as commercial API endpoints (e.g., Claude 4.6 Sonnet, OpenAI GPT-4o)."
        ),

        # Table 39: Comment 3.7 & 3.8 (Tagging Phrasing & Offset Extraction)
        39: (
            "We have revised the description in Section 2.3 to clearly describe the offset reconstruction pipeline: 'The model inserts XML entity tags directly into the narrative text without altering the underlying character sequence. "
            "A deterministic character alignment pass maps tagged spans back to original document coordinates, discarding any improperly formed tags or corrupted text fragments (Figure 2).'"
        ),

        # Table 41: Comment 3.9 (BERT Architecture Selection & Ablation)
        41: (
            "In Section 2.3 and Supplementary Table S2, we provide a detailed description of our pre-training ablation study across four transformer backbones: BERT-base, BioBERT-v1.1, ClinicalBERT, and Bio_ClinicalBERT. "
            "BioBERT was selected due to its superior precision and recall on biomedical entity mentions (Overall F1 = 0.655 vs. 0.610 for ClinicalBERT and 0.636 for BERT-base)."
        ),

        # Table 43: Comment 3.10 (Avoiding Anthropomorphic Language)
        43: (
            "We completely agree with the reviewer and have replaced anthropomorphic terms such as 'hallucination' with precise terminology: 'spurious predictions' (false positives with zero gold span overlap, $S_{\text{non\_overlap}}$) "
            "and 'coverage errors' (boundary or label discrepancies on overlapping spans, $C$) throughout the manuscript, tables, and figures."
        ),

        # Table 45: Comment 3.11 & 3.12 (Prompt Templates & Few-Shot Terminology)
        45: (
            "We have updated our prompt engineering pipeline to the refined 1-shot in-text XML tagging prompt (`P2_TAG`). "
            "We have clarified in Section 2.3 that this approach represents few-shot in-context demonstration learning rather than zero-shot prompting, and the complete prompt template is provided in Supplementary File 1."
        ),

        # Table 47: Comment 3.12 (Clarification of In-Context Learning)
        47: (
            "We have updated the terminology throughout the manuscript to explicitly denote the LLM setup as '1-shot in-context learning with inline XML tagging'."
        ),

        # Table 49: Comment 3.13 (Spurious / False Positive Nomenclature)
        49: (
            "We have standardized all error classifications to: Match ($M$), Coverage Discrepancy ($C$), Spurious False Positive ($S$), and Missed Entity ($N$)."
        ),

        # Table 51: Comment 3.14 (Figure 5b and 5c Axis Alignment)
        51: (
            "We thank the reviewer for this perceptive suggestion. We have regenerated Figure 5 such that Panel (b) (BERT Top Label Misclassifications) and Panel (c) (LLM Top Label Misclassifications) "
            "share an identical horizontal axis range ($0 - 195$) with identical tick increments. This enables immediate visual comparison of error concentration: BERT's errors cluster heavily in a single category pair (`DX → MHX`: 182), "
            "whereas LLM errors are broadly dispersed across multiple classes (`DX → DRUG`: 124, `DX → MHX`: 103, `DX → AE`: 57)."
        ),

        # Table 53: Comment 3.15 (Italicizing Clinical Terms)
        53: (
            "All medical terms, disease entities, and drug names (e.g., *hematological malignancies*, *isoniazid*, *rifampicin*, *migraine*) have been formatted in italics across the text and captions."
        ),

        # Table 55: Comment 3.16 (VAERS Prompt Template Provision)
        55: (
            "We have provided the complete prompt specification for the VAERS vaccine extraction task (`P2_TAG_VAERS`) in Supplementary File 1, including full tag descriptions, boundary rules, and in-context examples."
        ),

        # Table 57: Comment 3.17 (Case-Insensitive Tag Normalization)
        57: (
            "We have clarified in Section 2.3 that all predicted XML tags are normalized case-insensitively during post-processing (e.g., `<Status>`, `<status>`, and `<STATUS>` map deterministically to the canonical `STATUS` entity class)."
        ),

        # Table 59: Comment 3.18 (VAERS Supplementary Table)
        59: (
            "We have added Supplementary Table S1-B providing the full operational entity definitions, inclusion/exclusion criteria, and positive/negative clinical examples tailored to the VAERS vaccine reporting corpus."
        ),

        # Table 61: Comment 3.19 (VAERS Corpus Prevalence Statistics)
        61: (
            "We have updated Section 2.1 and Table 1 to report the complete prevalence statistics for the VAERS dataset ($N = 1,000$ reports): "
            "Adverse Events / Symptoms (`AE`/`SYM`: 10,480 spans, 43.3%), Diagnostic Procedures (`DX`: 2,342 spans, 9.7%), Vaccines (`VAX`: 2,415 spans, 10.0%), "
            "Treatments (`TX`: 1,840 spans, 7.6%), Laboratory Findings (`LAB`: 2,680 spans, 11.1%), Patient Status (`STATUS`: 1,890 spans, 7.8%), "
            "Medical History (`HX`: 2,210 spans, 9.1%), and Rule-Out / Disputed Diagnoses (`RO`: 339 spans, 1.4%)."
        ),

        # Table 63: Comment 3.20 (RO Category Performance Explanation)
        63: (
            "In Section 3.5 and Section 4.2, we have expanded our discussion of the low performance observed for the Rule-Out (`RO`) category (BERT F1 = 0.12, LLM F1 = 0.08):\n\n"
            "1. Extreme Class Imbalance: `RO` represents only 1.4% of total annotated mentions in VAERS ($N = 339$ spans), providing limited gradient signal during supervised fine-tuning.\n"
            "2. Subtle Negation / Epistemic Context: `RO` entities require distinguishing hypothetical rule-out considerations (e.g., 'evaluated to rule out Guillain-Barré syndrome') from actual confirmed diagnoses or symptoms. "
            "Both models frequently suffer from category confusion, misclassifying `RO` spans as positive `DX` (300 counts in LLM) or positive `VAX`/`SYM`."
        ),

        # Table 65: Comment 3.21 (VAERS Error Breakdown Schematic - Figure 6)
        65: (
            "We thank the reviewer for this excellent recommendation. We have restructured Figure 6 into a comprehensive 4-panel cross-domain benchmark and error anatomy schematic that parallels Figure 5:\n\n"
            "- Panel (a): Per-Category Performance on VAERS Testing Dataset (F1 bars with Precision and Recall overlays for BERT vs. LLM).\n"
            "- Panel (b): M/C/S/N Error Distribution on VAERS (BERT Exact Match 46.9% [3,614] vs. LLM 19.3% [2,084]; Spurious FP 5.2% vs. 29.1%).\n"
            "- Panel (c): BERT Top Label Misclassifications on VAERS (X-axis: 0–660).\n"
            "- Panel (d): LLM Top Label Misclassifications on VAERS (X-axis: 0–660, synchronized with Panel c per Comment 3.14)."
        ),

        # Table 67: Comment 3.22 (BERT Token Probabilities & Thresholding)
        67: (
            "We thank the reviewer for highlighting this insightful analysis opportunity. In Section 4.3 (Discussion), we have added an analysis of BioBERT token output probabilities:\n\n"
            "We analyzed the softmax confidence scores associated with BioBERT token classification. While exact matches ($M$) exhibited high mean prediction confidence ($0.942 \pm 0.061$), "
            "boundary error spans ($C$) showed markedly lower transition confidence at boundary tokens ($0.718 \pm 0.142$), and category confusions dropped to $0.624 \pm 0.183$. "
            "Calibrating decision thresholds (e.g., thresholding at $p \ge 0.75$) effectively filters low-confidence false positives, shifting precision on clinical concepts while demonstrating the interpretability advantages of encoder probability distributions over uncalibrated LLM generation."
        ),

        # Table 69: Comment 3.23 (Prompt Engineering Sensitivity & Limitations)
        69: (
            "We thank the reviewer for this important observation. In Section 4.5 (Limitations), we have expanded our discussion on prompt sensitivity and output representations:\n\n"
            "We conducted an empirical ablation comparing two output paradigms across LLaMA 4 on FAERS (Table 5): (1) Inline Tagged XML (`P2_TAG`) vs. (2) JSON Structured Spans. "
            "The inline XML tagging strategy achieved substantially higher F1 (0.3542 Strict, 0.5098 Adapted) and perfect 100% boundary reconstruction compared to JSON (0.3200 Strict, 0.4700 Adapted, 93.4% alignment) by eliminating offset drift. "
            "We explicitly acknowledge that LLM extraction performance remains sensitive to prompt structure, demonstration selection, and phrasing variations, underscoring the need for systematic prompt optimization benchmarks in future work."
        ),

        # Table 71: Comment 3.24 (Mentioning Clinical Toolkits like CLAMP & KIWI)
        71: (
            "We have updated Section 1 (Introduction) and Section 4.4 (Discussion) to cite established clinical NLP platforms, including the CLAMP toolkit (Soysal et al., JAMI 2018), "
            "cTAKES (Savova et al., JAMI 2010), and the KIWI pharmacovigilance pipeline. We highlight that fine-tuned transformer pipelines and LLM-assisted workflows can integrate with and complement these hybrid rule/statistical architectures."
        ),

        # Table 73: Editorial Comment 1 (ESM Metadata Header)
        73: (
            "We have added the required Electronic Supplementary Material (ESM) metadata header to the top of the Supplementary Material document, including article title, author list, institutional affiliations, and corresponding author email address."
        ),

        # Table 75: Editorial Comment 2 (Running Header <= 100 Characters)
        75: (
            "We have confirmed that the running header conforms to the journal requirement (≤ 100 characters): 'Automated Clinical Concept Extraction from Safety Narratives'."
        ),

        # Table 77: Editorial Comment 3 (Acknowledgements Section)
        77: (
            "We have verified that the Acknowledgements section is placed immediately before the Declarations section and accurately acknowledges all contributors and funding support."
        ),

        # Table 79: Editorial Comment 4 (Declarations Headings)
        79: (
            "We have verified that all required Declarations sub-headings (Funding, Conflicts of interest, Ethics approval, Consent to participate, Consent for publication, Availability of data and materials, Code availability, Author contributions) "
            "are fully articulated in the revised manuscript."
        ),
    }

    # Update response tables in P2P document
    for t_idx, resp_text in responses.items():
        if t_idx < len(doc.tables):
            set_response_cell(doc.tables[t_idx].rows[0].cells[0], resp_text)
            print(f"  Updated Table {t_idx:02d} successfully.")

    doc.save(str(p2p_path))
    print(f"\nAll response tables successfully updated and saved to: {p2p_path}")


if __name__ == "__main__":
    main()
