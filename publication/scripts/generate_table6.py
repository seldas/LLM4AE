#!/usr/bin/env python3
"""Generate publication-ready Table 6: BioBERT Optimization Stability and Random Seed Invariance.

Evaluates Supervised BioBERT optimization stability across 5 independent random initialization seeds
(seeds 42, 123, 456, 789, 1011) on:
1. FAERS 4-Fold LOO Benchmark (N = 829 Reports Total, 17 categories):
   - Each seed evaluated as full-corpus micro-average F1 (pooled across all 829 reports).
   - Pooled row reported as Mean ± SD across the 5 seeds (Strict: 0.5564 ± 0.0069, Adapted: 0.7420 ± 0.0061),
     strictly aligning with Table 5 Total row.
2. VAERS 10-Fold CV Benchmark (N = 1,000 Reports Total, 14 categories):
   - 10-Fold CV cross-fold Mean ± SD per seed.
   - Pooled row reported as Mean ± SD across the 5 seeds (Strict: 0.7009 ± 0.0016, Adapted: 0.8397 ± 0.0018).

Data sources:
- FAERS: publication/results/bert_runs_FAERS_LOO/raw.xlsx
- VAERS: publication/results/comparison_three_schemes/three_schemes_summary.xlsx

Outputs:
- publication/results/tables/table6_random_seed_invariance.md
- publication/manuscripts/Tables/table6.md
- publication/manuscripts/table6.md
- publication/results/tables/table6_random_seed_invariance.xlsx
- publication/results/tables/table6_data.json
- Updates Table 6 in publication/manuscripts/LLM4AE_rev1.docx
"""

from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
from typing import Any, Dict, List, Tuple

import pandas as pd
import docx


PROJECT_ROOT = Path(__file__).resolve().parent.parent
BERT_FAERS_RAW_PATH_DEFAULT = PROJECT_ROOT / "results" / "bert_runs_FAERS_LOO" / "raw.xlsx"
VAERS_SUMMARY_PATH_DEFAULT = PROJECT_ROOT / "results" / "comparison_three_schemes" / "three_schemes_summary.xlsx"
TABLES_DIR_DEFAULT = PROJECT_ROOT / "results" / "tables"
MANUSCRIPT_DIR_DEFAULT = PROJECT_ROOT / "manuscripts"
DOCX_PATH_DEFAULT = MANUSCRIPT_DIR_DEFAULT / "LLM4AE_rev1.docx"


def calculate_two_tier_metrics(df: pd.DataFrame) -> Dict[str, float]:
    """Compute Strict Exact-Match F1 and Adapted ADE-Eval F1 from raw aligned rows."""
    if df.empty:
        return {"strict_F1": 0.0, "ade_F1": 0.0, "strict_P": 0.0, "strict_R": 0.0, "ade_P": 0.0, "ade_R": 0.0}

    counts = df["match_type"].value_counts().to_dict()
    M = int(counts.get("M", 0))
    class_confusion_mask = (df["match_type"] == "C") & (df["error_subtype"] == "class_confusion") if "error_subtype" in df.columns else pd.Series(False, index=df.index)
    C_class = int(class_confusion_mask.sum())
    C_total = int(counts.get("C", 0))
    C_boundary = C_total - C_class
    S_non_overlap = int(counts.get("S", 0))
    N = int(counts.get("N", 0))

    # Scheme 3: Strict Exact-Match
    p3_den = M + C_total + S_non_overlap
    r3_den = M + C_total + N
    p3 = M / p3_den if p3_den > 0 else 0.0
    r3 = M / r3_den if r3_den > 0 else 0.0
    f3 = 2 * p3 * r3 / (p3 + r3) if (p3 + r3) > 0 else 0.0

    # Scheme 2: Adapted ADE-Eval
    m2 = M + 0.5 * C_total
    p2_den = M + C_total + 0.25 * S_non_overlap
    r2_den = M + C_total + N
    p2 = m2 / p2_den if p2_den > 0 else 0.0
    r2 = m2 / r2_den if r2_den > 0 else 0.0
    f2 = 2 * p2 * r2 / (p2 + r2) if (p2 + r2) > 0 else 0.0

    return {
        "strict_P": round(p3, 4),
        "strict_R": round(r3, 4),
        "strict_F1": round(f3, 4),
        "ade_P": round(p2, 4),
        "ade_R": round(r2, 4),
        "ade_F1": round(f2, 4),
    }


def compute_table6_data(faers_raw_path: Path, vaers_summary_path: Path) -> Tuple[List[dict], dict]:
    """Compute Table 6 metrics dynamically for FAERS (raw) and VAERS."""
    print(f"Loading FAERS raw records from: {faers_raw_path}")
    df_faers_raw = pd.read_excel(faers_raw_path, sheet_name="Raw_Results")
    faers_seeds = sorted(int(s) for s in df_faers_raw["seed"].dropna().unique())
    print(f"FAERS random seeds found: {faers_seeds}")

    table_rows = []
    audit_data: Dict[str, Any] = {
        "FAERS_4Fold_LOO": {},
        "VAERS_10Fold_CV": {},
    }

    # 1. FAERS Seeds (Micro-Average across all 829 reports for each seed)
    faers_strict_list, faers_ade_list = [], []
    for seed in faers_seeds:
        sub = df_faers_raw[df_faers_raw["seed"] == seed]
        met = calculate_two_tier_metrics(sub)
        faers_strict_list.append(met["strict_F1"])
        faers_ade_list.append(met["ade_F1"])

        table_rows.append({
            "Dataset & Evaluation Protocol": "FAERS (4-Fold LOO, N = 829)",
            "Random Seed": f"Seed {seed}",
            "Primary Tier: Strict Exact F1": f"{met['strict_F1']:.4f}",
            "Secondary Tier: Adapted ADE F1": f"{met['ade_F1']:.4f}",
            "_is_pooled": False,
            "_dataset": "FAERS",
        })
        audit_data["FAERS_4Fold_LOO"][f"Seed_{seed}"] = met

    f_s_series = pd.Series(faers_strict_list)
    f_a_series = pd.Series(faers_ade_list)
    f_s_mean, f_s_std = f_s_series.mean(), f_s_series.std()
    f_a_mean, f_a_std = f_a_series.mean(), f_a_series.std()

    table_rows.append({
        "Dataset & Evaluation Protocol": "FAERS (4-Fold LOO, Pooled)",
        "Random Seed": "Mean ± SD (5 Seeds)",
        "Primary Tier: Strict Exact F1": f"{f_s_mean:.4f} ± {f_s_std:.4f}",
        "Secondary Tier: Adapted ADE F1": f"{f_a_mean:.4f} ± {f_a_std:.4f}",
        "_is_pooled": True,
        "_dataset": "FAERS",
        "_strict_mean": f_s_mean,
        "_strict_std": f_s_std,
        "_ade_mean": f_a_mean,
        "_ade_std": f_a_std,
    })
    audit_data["FAERS_4Fold_LOO"]["Pooled_Summary"] = {
        "strict_F1_mean": round(float(f_s_mean), 4),
        "strict_F1_std": round(float(f_s_std), 4),
        "ade_F1_mean": round(float(f_a_mean), 4),
        "ade_F1_std": round(float(f_a_std), 4),
    }

    # 2. VAERS Seeds (from Seed_Ablation_VAERS)
    print(f"Loading VAERS summary from: {vaers_summary_path}")
    df_vaers_seeds = pd.read_excel(vaers_summary_path, sheet_name="Seed_Ablation_VAERS")

    for _, r in df_vaers_seeds.iterrows():
        seed_val = int(r["seed"])
        s_mean, s_std = float(r["Strict_F1_mean"]), float(r["Strict_F1_std"])
        a_mean, a_std = float(r["ADE_F1_mean"]), float(r["ADE_F1_std"])

        table_rows.append({
            "Dataset & Evaluation Protocol": "VAERS (10-Fold CV, N = 1,000)",
            "Random Seed": f"Seed {seed_val}",
            "Primary Tier: Strict Exact F1": f"{s_mean:.4f} ± {s_std:.4f}",
            "Secondary Tier: Adapted ADE F1": f"{a_mean:.4f} ± {a_std:.4f}",
            "_is_pooled": False,
            "_dataset": "VAERS",
        })
        audit_data["VAERS_10Fold_CV"][f"Seed_{seed_val}"] = {
            "strict_F1_mean": s_mean,
            "strict_F1_std": s_std,
            "ade_F1_mean": a_mean,
            "ade_F1_std": a_std,
        }

    v_strict_mean = float(df_vaers_seeds["Strict_F1_mean"].mean())
    v_strict_std = float(df_vaers_seeds["Strict_F1_mean"].std())
    v_ade_mean = float(df_vaers_seeds["ADE_F1_mean"].mean())
    v_ade_std = float(df_vaers_seeds["ADE_F1_mean"].std())

    table_rows.append({
        "Dataset & Evaluation Protocol": "VAERS (10-Fold CV, Pooled)",
        "Random Seed": "Mean ± SD (5 Seeds)",
        "Primary Tier: Strict Exact F1": f"{v_strict_mean:.4f} ± {v_strict_std:.4f}",
        "Secondary Tier: Adapted ADE F1": f"{v_ade_mean:.4f} ± {v_ade_std:.4f}",
        "_is_pooled": True,
        "_dataset": "VAERS",
        "_strict_mean": v_strict_mean,
        "_strict_std": v_strict_std,
        "_ade_mean": v_ade_mean,
        "_ade_std": v_ade_std,
    })
    audit_data["VAERS_10Fold_CV"]["Pooled_Summary"] = {
        "strict_F1_mean": round(v_strict_mean, 4),
        "strict_F1_std": round(v_strict_std, 4),
        "ade_F1_mean": round(v_ade_mean, 4),
        "ade_F1_std": round(v_ade_std, 4),
    }

    return table_rows, audit_data


def generate_markdown_content(table_rows: List[dict]) -> str:
    """Format Table 6 rows into clean markdown."""
    lines = [
        "# Table 6: BioBERT Optimization Stability and Performance Invariance Across Five Independent Random Initialization Seeds",
        "",
        "Evaluation of neural network optimization stability across 5 independent training runs (seeds 42, 123, 456, 789, 1011) for supervised BioBERT on the FAERS 4-fold LOO benchmark (20 total model runs, 17 categories) and VAERS 10-fold CV benchmark (50 total model runs).",
        "",
        "| Dataset & Evaluation Protocol | Random Seed | Primary Tier: Strict Exact F1 | Secondary Tier: Adapted ADE F1 |",
        "| :--- | :--- | :---: | :---: |",
    ]

    for r in table_rows:
        is_p = r["_is_pooled"]
        ds_name = f"**{r['Dataset & Evaluation Protocol']}**" if is_p else r["Dataset & Evaluation Protocol"]
        seed_name = f"**{r['Random Seed']}**" if is_p else r["Random Seed"]
        f1_s = f"**{r['Primary Tier: Strict Exact F1']}**" if is_p else r["Primary Tier: Strict Exact F1"]
        f1_a = f"**{r['Secondary Tier: Adapted ADE F1']}**" if is_p else r["Secondary Tier: Adapted ADE F1"]
        lines.append(f"| {ds_name} | {seed_name} | {f1_s} | {f1_a} |")

    faers_pooled = [r for r in table_rows if r["_is_pooled"] and r["_dataset"] == "FAERS"][0]
    vaers_pooled = [r for r in table_rows if r["_is_pooled"] and r["_dataset"] == "VAERS"][0]

    lines.extend([
        "",
        "---",
        "",
        "### Footnotes & Methodological Notes:",
        "1. **FAERS Protocol:** 4-fold Leave-One-Drug-Event-Pair-Out cross-validation evaluated across all 17 clinical concept categories. For each seed, results reflect full-corpus micro-average F1 pooled across all 829 reports.",
        "2. **VAERS Protocol:** 10-fold cross-validation on the 1,000 VAERS reports. For each seed, Mean $\\pm$ SD represents cross-fold variation across the 10 test partitions.",
        f"3. **Pooled Invariance:** The pooled summary represents the Mean $\\pm$ SD across the 5 independent random initialization seeds, demonstrating minimal stochastic variation ($SD = {faers_pooled['_strict_std']:.4f}$ on FAERS, $SD = {vaers_pooled['_strict_std']:.4f}$ on VAERS).",
        "",
    ])
    return "\n".join(lines)


def update_docx_manuscript(docx_path: Path, table_rows: List[dict]) -> None:
    """Update Table 6 (Table index 5 in docx) and associated text in LLM4AE_rev1.docx."""
    if not docx_path.exists():
        print(f"Warning: docx manuscript not found at {docx_path}")
        return

    doc = docx.Document(str(docx_path))

    # Locate Table 6 (Table 5 in doc.tables index)
    target_table = None
    for t in doc.tables:
        header_text = " ".join(c.text for c in t.rows[0].cells).lower()
        if "random seed" in header_text and "strict exact" in header_text:
            target_table = t
            break

    if target_table is not None:
        print(f"Found Table 6 in {docx_path.name}; updating table cells...")
        for row_idx, r_data in enumerate(table_rows, start=1):
            if row_idx < len(target_table.rows):
                row_cells = target_table.rows[row_idx].cells
                if r_data["_is_pooled"]:
                    row_cells[0].text = r_data["_dataset"]
                    row_cells[1].text = "Average"
                else:
                    row_cells[1].text = r_data["Random Seed"]
                row_cells[2].text = r_data["Primary Tier: Strict Exact F1"]
                row_cells[3].text = r_data["Secondary Tier: Adapted ADE F1"]

    # Update paragraph P121 text
    faers_pooled = [r for r in table_rows if r["_is_pooled"] and r["_dataset"] == "FAERS"][0]
    vaers_pooled = [r for r in table_rows if r["_is_pooled"] and r["_dataset"] == "VAERS"][0]

    for p in doc.paragraphs:
        if "Table 6" in p.text and "optimization stability" in p.text.lower() and "variance" in p.text.lower():
            p.text = (
                "To rigorously verify neural network optimization stability, we conducted 5 independent training runs "
                "using random initialization seeds (42, 123, 456, 789, 1011) across both FAERS and VAERS. As summarized in "
                f"Table 6, cross-runs variance was exceptionally low across both datasets. On FAERS, strict F1 averaged "
                f"{faers_pooled['Primary Tier: Strict Exact F1']} ({faers_pooled['Secondary Tier: Adapted ADE F1']} Adapted) across the 5 seeds. "
                f"On VAERS, strict F1 averaged {vaers_pooled['Primary Tier: Strict Exact F1']} ({vaers_pooled['Secondary Tier: Adapted ADE F1']} Adapted) "
                "across the 5 seeds. These findings confirm that BERT convergence is robust to weight initialization."
            )
            break

    try:
        doc.save(str(docx_path))
        print(f"Successfully updated Table 6 and narrative in: {docx_path}")
    except PermissionError:
        alt_path = docx_path.with_name(f"{docx_path.stem}_table6_updated{docx_path.suffix}")
        doc.save(str(alt_path))
        print(
            f"Notice: '{docx_path.name}' is currently open in Microsoft Word (locked).\n"
            f"Saved updated version to: '{alt_path}'."
        )
    except Exception as e:
        print(f"Warning: Could not save updated docx: {e}")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--faers-raw", type=Path, default=BERT_FAERS_RAW_PATH_DEFAULT, help="Path to bert_runs_FAERS_LOO/raw.xlsx")
    parser.add_argument("--vaers-summary", type=Path, default=VAERS_SUMMARY_PATH_DEFAULT, help="Path to three_schemes_summary.xlsx")
    parser.add_argument("--tables-dir", type=Path, default=TABLES_DIR_DEFAULT, help="Output directory for table artifacts")
    parser.add_argument("--manuscript-dir", type=Path, default=MANUSCRIPT_DIR_DEFAULT, help="Path to manuscripts directory")
    parser.add_argument("--docx-path", type=Path, default=DOCX_PATH_DEFAULT, help="Path to LLM4AE_rev1.docx")
    args = parser.parse_args()

    args.tables_dir.mkdir(parents=True, exist_ok=True)
    args.manuscript_dir.mkdir(parents=True, exist_ok=True)

    table_rows, audit_data = compute_table6_data(args.faers_raw, args.vaers_summary)

    # 1. Print Summary
    print("\n=== Table 6: BioBERT Optimization Stability Across Random Seeds ===")
    for r in table_rows:
        print(f"{r['Dataset & Evaluation Protocol']:30s} | {r['Random Seed']:20s} | Strict: {r['Primary Tier: Strict Exact F1']:17s} | Adapted: {r['Secondary Tier: Adapted ADE F1']:17s}")

    # 2. Markdown Files
    md_content = generate_markdown_content(table_rows)
    md_paths = [
        args.tables_dir / "table6_random_seed_invariance.md",
        args.manuscript_dir / "Tables" / "table6.md",
        args.manuscript_dir / "table6.md",
    ]
    for p in md_paths:
        p.parent.mkdir(parents=True, exist_ok=True)
        with p.open("w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"Saved markdown: {p}")

    # 3. Excel Workbook with ESM Cover Sheet
    esm_cover = pd.DataFrame([
        {"Metadata Field": "Article Title", "Value": "Benchmarking Fine-Tuned Encoders and Instruction-Tuned Large Language Models for Adverse Event Clinical Concept Extraction from Spontaneous Reporting Narratives"},
        {"Metadata Field": "Journal", "Value": "Drug Safety"},
        {"Metadata Field": "Table Identifier", "Value": "Table 6: BioBERT Optimization Stability and Random Seed Invariance"},
        {"Metadata Field": "Datasets", "Value": "FAERS 4-Fold LOO (N = 829, 17 Categories) & VAERS 10-Fold CV (N = 1,000 Reports)"},
        {"Metadata Field": "Random Seeds", "Value": "5 Independent Initialization Seeds (42, 123, 456, 789, 1011)"},
        {"Metadata Field": "Generated By", "Value": "publication/scripts/generate_table6.py (Computed dynamically from raw.xlsx)"}
    ])

    export_df = pd.DataFrame([
        {
            "Dataset & Evaluation Protocol": r["Dataset & Evaluation Protocol"],
            "Random Seed": r["Random Seed"],
            "Primary Tier: Strict Exact F1": r["Primary Tier: Strict Exact F1"],
            "Secondary Tier: Adapted ADE F1": r["Secondary Tier: Adapted ADE F1"],
        }
        for r in table_rows
    ])

    out_excel = args.tables_dir / "table6_random_seed_invariance.xlsx"
    with pd.ExcelWriter(out_excel, engine="openpyxl") as writer:
        esm_cover.to_excel(writer, sheet_name="ESM_Cover_Sheet", index=False)
        export_df.to_excel(writer, sheet_name="Table_6_Seed_Invariance", index=False)
    print(f"Saved excel: {out_excel}")

    # 4. JSON Audit Data
    out_json = args.tables_dir / "table6_data.json"
    with out_json.open("w", encoding="utf-8") as f:
        json.dump(audit_data, f, indent=2)
    print(f"Saved audit json: {out_json}")

    # 5. Update Word Manuscript Docx
    update_docx_manuscript(args.docx_path, table_rows)


if __name__ == "__main__":
    main()
