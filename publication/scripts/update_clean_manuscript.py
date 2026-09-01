#!/usr/bin/env python3
"""
update_clean_manuscript.py

Synchronizes LLM4AE_rev1_clean_updated.docx with:
1. High-resolution Figures 2-6 (direct ZIP media replacement).
2. Clean, synchronized publication tables evaluated across the full 17 categories:
   - Table 1 (Section 2.1): Descriptive statistics of the annotated corpora (FAERS & VAERS)
   - Table 2 (Section 3.1): FAERS annotations, categorized by Human, ETHER and LLM
   - Table 3 (Section 3.3): Master Performance Benchmark on FAERS (17 Categories: BioBERT, ClinicalBERT, LLaMA4 Tagged/JSON, Sonnet, ETHER)
   - Table 4 (Section 3.5): Master Performance Benchmark on VAERS (14 Categories: BioBERT, LLaMA 4)
   - Table 5 (Section 3.6): Leave-One-Drug-Event-Pair-Out Performance across 4 Case Series on FAERS (17 Categories)
   - Table 6 (Section 3.6): BioBERT Random Seed Invariance Across 5 Seeds on FAERS (4-Fold LOO, 17 Categories) and VAERS (10-Fold CV, 14 Categories)
   - Table 7 (Section 3.6): LLM Output Format Paradigm Comparison (Inline Tagged XML vs JSON Schema, 17 Categories)
   - Table 8 (Section 3.6): Pretrained Transformer Encoder Architecture Ablation on VAERS (BioBERT vs Bio_ClinicalBERT vs BERT-Base vs ClinicalBERT, 5 Seeds)
3. Results text cleaned so all numbers across Table 3, Table 5, Table 6, Table 7, and Table 8 are 100% consistent with 17-category evaluation.
"""

from __future__ import annotations

import io
import shutil
import zipfile
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def replace_media_images_in_memory(docx_path: Path, img_map: dict[str, Path]):
    """Replaces images inside the docx zip archive using an in-memory buffer."""
    with open(docx_path, 'rb') as f:
        in_buf = io.BytesIO(f.read())
    
    out_buf = io.BytesIO()
    with zipfile.ZipFile(in_buf, 'r') as zin, zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in img_map and img_map[item.filename].exists():
                print(f"  Replacing {item.filename} with {img_map[item.filename].name} ({img_map[item.filename].stat().st_size} bytes)...")
                with open(img_map[item.filename], "rb") as fimg:
                    zout.writestr(item.filename, fimg.read())
            else:
                zout.writestr(item, zin.read(item.filename))
    
    with open(docx_path, 'wb') as f:
        f.write(out_buf.getvalue())
    print("Media replacement complete.")


def create_styled_table(doc, data: list[list[str]], col_widths: list[float] | None = None, is_header_multiline: int = 1):
    num_rows = len(data)
    num_cols = len(data[0])
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="8" w:space="0" w:color="333333"/>'
            f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="333333"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx]
        is_hdr = (r_idx < is_header_multiline)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            tcPr = cell._element.get_or_add_tcPr()
            if is_hdr:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F4F8"/>')
            else:
                bg = "FFFFFF" if r_idx % 2 == 1 else "FBFBFB"
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            tcPr.append(shd)

            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="60" w:type="dxa"/>'
                f'<w:bottom w:w="60" w:type="dxa"/>'
                f'<w:left w:w="90" w:type="dxa"/>'
                f'<w:right w:w="90" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(tcMar)

            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (is_hdr or c_idx > 0) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            run = p.add_run(val)
            run.bold = is_hdr or (c_idx == 0) or ("OVERALL" in val) or ("Mean" in val) or ("BioBERT" in val) or ("Pooled" in val)
            run.font.name = "Arial"
            run.font.size = Pt(8.0)
            run.font.color.rgb = RGBColor(0x11, 0x11, 0x11) if is_hdr else RGBColor(0x22, 0x22, 0x22)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for r in table.rows:
            for c_idx, w in enumerate(col_widths):
                if c_idx < len(r.cells):
                    r.cells[c_idx].width = Inches(w)

    return table


def replace_table_in_place(old_table, new_table):
    parent = old_table._tbl.getparent()
    parent.insert(parent.index(old_table._tbl), new_table._tbl)
    parent.remove(old_table._tbl)


def remove_table(table):
    parent = table._tbl.getparent()
    parent.remove(table._tbl)


def main():
    repo_root = Path(__file__).resolve().parent.parent.parent
    manuscript_dir = repo_root / "publication" / "manuscripts"
    src_docx_path = manuscript_dir / "LLM4AE_rev1_clean_updated.docx"
    dst_docx_path = manuscript_dir / "LLM4AE_rev1_clean_updated.docx"

    print(f"Loading base document from {src_docx_path}...")
    doc = docx.Document(str(src_docx_path))

    # --- TABLE 3: MASTER BENCHMARK ON FAERS (Section 3.3, 17 Categories) ---
    t3_data = [
        ["Model & Configuration", "Primary Tier: Strict Exact F1", "Secondary Tier: Adapted ADE F1"],
        ["BioBERT (4-Fold LOO)", "0.5685 ± 0.0080", "0.7463 ± 0.0076"],
        ["LLaMA 4 (1-shot, Tagged P2_TAG)", "0.4043", "0.6249"],
        ["Claude 4.6 Sonnet (1-shot, Tagged)", "0.4667", "0.6443"],
        ["ETHER (Baseline, used=Yes)", "0.1147", "0.2447"]
    ]
    t3_styled = create_styled_table(doc, t3_data, col_widths=[2.8, 1.8, 1.8])

    # --- TABLE 4: MASTER BENCHMARK ON VAERS (Section 3.5, 14 Categories) ---
    t4_data = [
        ["Model & Configuration", "Primary Tier: Strict Exact F1", "Secondary Tier: Adapted ADE F1"],
        ["BioBERT (10-Fold CV)", "0.6594 ± 0.0196", "0.7848 ± 0.0127"],
        ["LLaMA 4 (1-shot, Tagged P2_TAG_VAERS)", "0.2364", "0.4766"]
    ]
    t4_styled = create_styled_table(doc, t4_data, col_widths=[2.8, 1.8, 1.8])

    # --- TABLE 5: LEAVE-ONE-DRUG-EVENT-PAIR-OUT (Section 3.6, 17 Categories) ---
    t5_data = [
        ["Drug–Event Case Series", "Validation Cohort Size", "Primary Tier: Strict Exact F1", "Secondary Tier: Adapted ADE F1"],
        ["Azacitidine – QT Prolongation", "N = 200 reports", "0.6002 ± 0.0114", "0.7733 ± 0.0092"],
        ["Baricitinib – Hypersensitivity", "N = 200 reports", "0.6367 ± 0.0201", "0.7751 ± 0.0128"],
        ["Tramadol – Hypoglycemia", "N = 229 reports", "0.5289 ± 0.0093", "0.7242 ± 0.0036"],
        ["Erenumab – Stroke", "N = 200 reports", "0.5084 ± 0.0122", "0.7126 ± 0.0077"],
        ["Macro-Average (All 4 Folds)", "N = 829 reports total", "0.5685 ± 0.0080", "0.7463 ± 0.0076"]
    ]
    t5_styled = create_styled_table(doc, t5_data, col_widths=[2.4, 1.5, 1.6, 1.6])

    # --- TABLE 6: RANDOM SEED INVARIANCE TABLE (Section 3.6, 17 FAERS / 14 VAERS Categories) ---
    t6_data = [
        ["Dataset & Evaluation Protocol", "Random Seed", "Primary Tier: Strict Exact F1", "Secondary Tier: Adapted ADE F1"],
        ["FAERS (4-Fold LOO, N = 829)", "Seed 42", "0.5582 ± 0.0649", "0.7431 ± 0.0295"],
        ["FAERS (4-Fold LOO, N = 829)", "Seed 123", "0.5652 ± 0.0509", "0.7402 ± 0.0293"],
        ["FAERS (4-Fold LOO, N = 829)", "Seed 456", "0.5783 ± 0.0697", "0.7586 ± 0.0355"],
        ["FAERS (4-Fold LOO, N = 829)", "Seed 789", "0.5662 ± 0.0476", "0.7413 ± 0.0324"],
        ["FAERS (4-Fold LOO, N = 829)", "Seed 1011", "0.5748 ± 0.0694", "0.7484 ± 0.0371"],
        ["FAERS (4-Fold LOO, Pooled)", "Mean ± SD (5 Seeds)", "0.5685 ± 0.0080", "0.7463 ± 0.0076"],
        ["VAERS (10-Fold CV, N = 1,000)", "Seed 42", "0.6594 ± 0.0196", "0.7848 ± 0.0127"],
        ["VAERS (10-Fold CV, N = 1,000)", "Seed 123", "0.6601 ± 0.0175", "0.7891 ± 0.0104"],
        ["VAERS (10-Fold CV, N = 1,000)", "Seed 456", "0.6593 ± 0.0218", "0.7883 ± 0.0138"],
        ["VAERS (10-Fold CV, N = 1,000)", "Seed 789", "0.6615 ± 0.0159", "0.7907 ± 0.0095"],
        ["VAERS (10-Fold CV, N = 1,000)", "Seed 1011", "0.6574 ± 0.0169", "0.7879 ± 0.0106"],
        ["VAERS (10-Fold CV, Pooled)", "Mean ± SD (5 Seeds)", "0.6595 ± 0.0015", "0.7882 ± 0.0022"]
    ]
    t6_styled = create_styled_table(doc, t6_data, col_widths=[2.4, 1.6, 1.8, 1.8])

    # --- TABLE 7: OUTPUT FORMAT PARADIGM COMPARISON (Section 3.6, 17 Categories) ---
    t7_data = [
        ["Model", "Prompt Strategy & Output Paradigm", "Strict Exact-Match F1", "Adapted ADE-Eval F1", "Boundary Alignment Success"],
        ["LLaMA 4 (1-shot)", "Inline Tagged XML (P2_TAG)", "0.4043", "0.6249", "100.0%"],
        ["LLaMA 4 (1-shot)", "JSON Schema (Structured Span Offsets)", "0.4071", "0.5995", "93.4%"],
        ["Claude 4.6 Sonnet (1-shot)", "Inline Tagged XML (P2_TAG)", "0.4667", "0.6443", "100.0%"]
    ]
    t7_styled = create_styled_table(doc, t7_data, col_widths=[1.6, 2.3, 1.2, 1.2, 1.2])

    # --- TABLE 8: PRETRAINED ENCODER ABLATION (Section 3.6, 4 BERT Variants x 5 Seeds) ---
    t8_data = [
        ["Model Architecture", "Validation F1 (Mean ± SD)", "Validation Precision", "Validation Recall", "Clinical Score"],
        ["BioBERT v1.1", "0.8471 ± 0.0058", "0.8666 ± 0.0048", "0.8285 ± 0.0089", "0.8500 ± 0.0071"],
        ["Bio_ClinicalBERT", "0.8433 ± 0.0070", "0.8610 ± 0.0078", "0.8264 ± 0.0104", "0.8440 ± 0.0055"],
        ["BERT-Base", "0.8382 ± 0.0047", "0.8596 ± 0.0055", "0.8179 ± 0.0061", "0.8420 ± 0.0045"],
        ["ClinicalBERT", "0.8369 ± 0.0086", "0.8615 ± 0.0106", "0.8140 ± 0.0167", "0.8400 ± 0.0071"]
    ]
    t8_styled = create_styled_table(doc, t8_data, col_widths=[1.8, 1.4, 1.4, 1.4, 1.2])

    # Update in-place Tables 2 through 7 (which correspond to Tables 3-8 in manuscript numbering)
    replace_table_in_place(doc.tables[2], t3_styled)
    replace_table_in_place(doc.tables[3], t4_styled)
    replace_table_in_place(doc.tables[4], t5_styled)
    replace_table_in_place(doc.tables[5], t6_styled)
    replace_table_in_place(doc.tables[6], t7_styled)
    replace_table_in_place(doc.tables[7], t8_styled)

    # Clean up duplicate table if len(doc.tables) > 8
    if len(doc.tables) > 8:
        for extra_idx in range(len(doc.tables) - 1, 7, -1):
            remove_table(doc.tables[extra_idx])

    # Clean Paragraphs & Captions
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        
        # Section 3.3 (FAERS Benchmark text)
        if txt.startswith("We trained and evaluated a supervised BioBERT"):
            p.text = (
                "We trained and evaluated a supervised BioBERT named entity recognition (NER) model "
                "using a 4-fold Leave-One-Drug-Event-Pair-Out cross-validation design on the 829 FAERS narratives across all 17 clinical concept categories. "
                "Table 3 summarizes the overall benchmark comparison across model families, input paradigms, and evaluation tiers. "
                "BioBERT achieved a strict exact-match F1 of 0.5685 ± 0.0080 and an adapted ADE-Eval F1 of 0.7463 ± 0.0076, "
                "which substantially outperformed few-shot LLaMA 4 (0.4043 Strict, 0.6249 Adapted), "
                "Claude 4.6 Sonnet (0.4667 Strict, 0.6443 Adapted), and baseline ETHER system (0.1147 Strict, 0.2447 Adapted)."
            )
        
        # Section 3.3 Figure 4 caption
        elif "Fig. 4" in txt or "Figure 4" in txt:
            if "Caption:" in txt or txt.startswith("Figure 4.") or txt.startswith("Fig. 4"):
                p.text = "Figure 4. Comparative Concept Extraction Performance Across All 17 Clinical Concept Categories on the FAERS Benchmark Corpus (N = 829 Reports) for Fine-Tuned BioBERT (Blue Diamond / Bar), Claude 4.6 Sonnet (Red Circle / Bar), and LLaMA 4 (Pink Square / Bar)."

        # Section 3.4 (Error Analysis text & caption)
        elif txt.startswith("Next, we conducted a detailed error analysis"):
            p.text = (
                "Next, we conducted a detailed error analysis on the complete FAERS benchmark corpus (N = 829 reports) "
                "to characterize how the supervised BERT model and the few-shot LLM differ in their failure modes. "
                "Using character-level span alignment between model predictions and curated reference annotations across all 17 clinical concept categories, "
                "each predicted entity was categorized as an exact match (M), coverage error (C, overlapping a reference span but with mismatched start/end boundaries "
                "or a conflated category label), or spurious prediction / hallucination (S, non-overlapping span with no corresponding reference entity). "
                "Unpredicted reference entities were counted as misses (N)."
            )
        elif txt.startswith("Figure 5a shows that BERT produces"):
            p.text = (
                "Figure 5a shows that BERT produces substantially more exact matches and fewer boundary, conflation, and hallucination errors "
                "than the LLM across the full FAERS benchmark corpus (N = 829 reports). For BERT, 58.0% of spans are exact matches (26,891), "
                "compared with 31.6% (17,396) for the LLM. The LLM, in contrast, exhibits more than double the proportion of coverage errors "
                "(29.7% vs 14.0%; 16,358 vs 6,485 spans), where a predicted span overlaps the reference span but has inaccurate boundaries or conflated labels. "
                "The LLM also produces significantly more spurious spans (25.7% vs 15.5%; 14,185 vs 7,195), reflecting hallucinated entities or over-segmentation. "
                "Missed spans (N) show comparable proportions between both paradigms (13.0% vs 12.5%; 7,156 vs 5,793 spans), which is consistent with the recall "
                "patterns observed in Figure 4, while the much larger excess of coverage and spurious errors for the LLM reflects its tendency to over-segment or "
                "hallucinate entities rather than leaving spans unannotated."
            )
        elif txt.startswith("The top confusion pairs for both systems"):
            p.text = (
                "The top confusion pairs for both systems (Figure 5b and 5c) demonstrate that errors cluster within clinically related concept types "
                "rather than dispersing randomly. For BERT, misclassifications primarily involve context-dependent clinical boundaries, notably medical history "
                "conflated with therapeutic indications (MHx → Indication: 345), adverse events confused with laboratory findings (AE → Lab: 191), and bidirectional "
                "medical history versus AE boundary ambiguities (MHx → AE: 172; AE → MHx: 168), typically occurring when chronic comorbidities and acute symptoms "
                "are documented within the same sentence. For the LLM, the single largest error mode is the misclassification of concomitant medications as primary "
                "suspect drugs (cDrug → sDrug: 2,737), followed by treatment interventions labeled as suspect drugs (Treatment → sDrug: 847). Additional prominent "
                "LLM confusions include medical history categorized as current diagnoses (MHx → Dx: 364) or adverse events (MHx → AE: 342), laboratory values labeled "
                "as dosages (Lab → Dose: 328), and laboratory findings classified as adverse events (Lab → AE: 230). While BERT errors are concentrated in a tight set "
                "of adjacent clinical categories, the LLM exhibits widespread role conflation across pharmacological and historical categories."
            )
        elif txt.startswith("We observed a low performance in annotating DX"):
            p.text = (
                "We applied word-cloud visualizations on the primary LLM misclassification categories to clarify these failure modes at the entity level. "
                "For concomitant and treatment spans misclassified as suspect drug (sDrug) by the LLM (Figure 5d), the most prominent tokens were common co-administered "
                "medications and supportive therapies, including methylprednisolone, methotrexate, prednisone, esomeprazole, folic acid, and dexamethasone. "
                "When multiple therapeutic agents appeared in complex clinical narratives, the few-shot LLM frequently assigned the primary sDrug label indiscriminately "
                "across all mentioned medications rather than recognizing their secondary concomitant or supportive roles. In contrast, clinical spans exhibiting "
                "medical history and diagnosis/AE confusions (Figure 5e) were dominated by prevalent chronic conditions and comorbidities, such as hypertension, "
                "myelodysplastic syndrome, pneumonia, rheumatoid arthritis, osteoporosis, and recurrent infections. This reflects a systematic tendency of the LLM "
                "to conflate chronological baseline disease history with active acute adverse event episodes."
            )
        elif txt.startswith("Overall, the LLM") and "few-shot annotations" in txt:
            p.text = (
                "Overall, the few-shot LLM annotations are prone to spurious span hallucinations, boundary drift, and broad drug role conflation "
                "(cDrug/Treatment → sDrug), which substantially depresses precision and necessitates extensive post-processing for downstream pharmacovigilance pipelines. "
                "Supervised BERT models, while strictly preserving schema conformity and entity boundaries, require targeted contextual features to resolve fine-grained "
                "distinctions between baseline medical history, therapeutic indications, and acute adverse event manifestations."
            )
        elif "Fig. 5" in txt or "Figure 5" in txt:
            if "Caption:" in txt or txt.startswith("Figure 5.") or txt.startswith("Fig. 5"):
                p.text = (
                    "Figure 5. Error Anatomy and Extraction Discrepancy Breakdown for Supervised BERT versus Few-Shot LLM (LLaMA 4) "
                    "on the Full FAERS Benchmark Corpus (N = 829 Reports). (a) Overall M/C/S/N error distribution (M: exact match; C: coverage error; S: spurious prediction; N: miss). "
                    "(b) Top label misclassifications for BERT. (c) Top label misclassifications for LLM. "
                    "(d) Typical concomitant drug and treatment terms misclassified as suspect drug (sDrug) by LLM. "
                    "(e) Typical clinical terms exhibiting medical history (MHx) versus diagnosis (Dx) or adverse event (AE) confusions by LLM."
                )

        # Section 3.5 (VAERS Benchmark text & caption)
        elif txt.startswith("Category-level analysis (Figure 6a) confirmed strong supervised encoder advantages"):
            p.text = (
                "Category-level analysis (Figure 6a) across all 14 clinical and contextual concept categories confirmed strong supervised encoder advantages "
                "across vaccines (VAX: 0.86 vs 0.60), treatments (Tx: 0.84 vs 0.40), patient status (STATUS: 0.71 vs 0.30), medical history (MHx: 0.75 vs 0.24), "
                "confirmed adverse event diagnoses (sDx: 0.69 vs 0.08), provisional diagnoses (pDx: 0.62 vs 0.07), symptoms (SYM: 0.63 vs 0.22), "
                "laboratory findings (Lab: 0.59 vs 0.16), and family history (FHx: 0.35 vs 0.10). "
                "Error anatomy (Figure 6b–d) on the complete VAERS corpus (N = 1,000 reports) revealed that BioBERT achieved a high proportion "
                "of exact matches (57.7% vs 13.7%; 15,068 vs 10,154 spans), while LLM few-shot predictions suffered from high spurious hallucinations "
                "(45.2% vs 18.6%; 33,421 vs 4,846 spans) and extensive missed entities (30.7% vs 17.1%; 22,711 vs 4,459 spans). "
                "Confusion breakdowns (Figure 6c & 6d) demonstrated that both models experienced diagnostic granularity conflations "
                "(sDx → DX: 4,902 for BERT, 3,116 for LLM; and SYM → DX: 2,782 for BERT, 2,999 for LLM)."
            )
        elif "Fig. 6" in txt or "Figure 6" in txt:
            if "Caption:" in txt or txt.startswith("Figure 6.") or txt.startswith("Fig. 6"):
                p.text = (
                    "Figure 6. Cross-Domain Annotation Benchmark and Error Anatomy on the Full VAERS Vaccine Safety Corpus Across 14 Concept Categories (N = 1,000 Reports). "
                    "(a) Per-category performance across all 14 clinical and contextual concept categories and overall micro-average for BioBERT (Blue) vs. LLaMA 4 (Pink-Coral), "
                    "showing F1 bars and precision/recall trajectories. (b) M/C/S/N error distribution for BioBERT vs. LLaMA 4 across the complete corpus. "
                    "(c) BioBERT top label misclassifications (X-axis: 0–5,400). (d) LLaMA 4 top label misclassifications (X-axis: 0–5,400, aligned with panel c)."
                )
        elif txt.startswith("Table 4.") and "VAERS" in txt:
            p.text = "Table 4. Master Performance Benchmark on the VAERS Dataset (N = 1,000 Reports)."

        # Section 3.6 Leave-One-Out text & caption
        elif txt.startswith("We applied a repeated Leave-One-Drug-Event-Pair-Out cross-validation protocol"):
            p.text = (
                "We applied a repeated Leave-One-Drug-Event-Pair-Out cross-validation protocol on the FAERS corpus "
                "to evaluate model generalization across distinct therapeutic contexts on all 17 clinical concept categories. For each of the 4 curated drug-event cohorts, "
                "the model was trained on the remaining 3 case series and evaluated on the held-out series. "
                "As shown in Table 5, strict exact-match F1 was 0.6002 ± 0.0114 for azacitidine–QT prolongation (N = 200), "
                "0.6367 ± 0.0201 for baricitinib–hypersensitivity (N = 200), 0.5289 ± 0.0093 for tramadol–hypoglycemia (N = 229), "
                "and 0.5084 ± 0.0122 for erenumab–stroke (N = 200), yielding a macro-average F1 of 0.5685 ± 0.0080 across folds "
                "(0.7463 ± 0.0076 Adapted F1). The between-series performance variance exceeded within-fold variation, "
                "demonstrating that narrative complexity and therapeutic vocabulary differences contribute more variation than stochastic model initialization."
            )
        elif txt.startswith("Table 5.") and "Leave-One" in txt:
            p.text = "Table 5. Leave-One-Drug-Event-Pair-Out Cross-Validation Performance Across Four FAERS Case Series (N = 829 Reports Total)."

        # Section 3.6 Random Seed section paragraph & caption
        elif txt.startswith("To rigorously verify neural network optimization stability"):
            p.text = (
                "To rigorously verify neural network optimization stability, we conducted 5 independent training runs "
                "using random initialization seeds (42, 123, 456, 789, 1011) across both FAERS 4-fold LOO (20 total model runs, 17 categories) "
                "and VAERS 10-fold CV (50 total model runs, 14 categories). As summarized in Table 6, cross-seed variance was exceptionally low across both datasets. "
                "On FAERS, strict F1 averaged 0.5685 ± 0.0080 (0.7463 ± 0.0076 Adapted F1) across the 5 seeds. "
                "On VAERS, strict F1 averaged 0.6595 ± 0.0015 (0.7882 ± 0.0022 Adapted F1) across the 5 seeds. "
                "These findings confirm that supervised BioBERT convergence is remarkably robust to stochastic weight initialization, "
                "with between-series clinical diversity accounting for substantially greater performance variance than random initialization."
            )
        elif txt.startswith("Table 6.") and ("BioBERT Optimization" in txt or "Random Initialization" in txt or "Invariance" in txt):
            p.text = "Table 6. BioBERT Optimization Stability and Performance Invariance Across Five Independent Random Initialization Seeds on FAERS (4-Fold LOO) and VAERS (10-Fold CV)."

        # Section 3.6 Output Format paragraph & caption
        elif txt.startswith("We investigated whether the output representation paradigm"):
            p.text = (
                "We investigated whether the output representation paradigm impacts generative LLM extraction fidelity. "
                "We benchmarked two structured output paradigms on FAERS (Table 7): (1) Inline Tagged XML (P2_TAG), where the LLM embeds XML tags directly into the narrative text, "
                "versus (2) JSON Schema, where the LLM produces a structured JSON array of extracted entity spans and character offsets across the 17 clinical concept categories. "
                "The inline XML tagging approach achieved higher adapted ADE F1 (0.6249 vs. 0.5995) and higher recall (0.5796 vs. 0.5232) with a 100.0% boundary alignment reconstruction rate, "
                "whereas JSON generation suppressed non-overlapping spurious entities by 25.9% (yielding higher strict precision, 0.3785 vs. 0.3470) but suffered from a 6.6% misalignment rate "
                "due to character offset hallucination and coordinate drift over long narratives."
            )
        elif txt.startswith("Table 7.") and ("Output Format" in txt or "Impact" in txt):
            p.text = "Table 7. Impact of LLM Output Format Paradigm (Inline Tagged XML vs. Structured JSON Schema Offsets) on Entity Extraction and Offset Alignment."

        # Section 3.6 BERT Encoder Architecture Ablation paragraph & caption
        elif txt.startswith("To investigate the impact of underlying pretraining corpora"):
            p.text = (
                "To investigate the impact of underlying pretraining corpora on pharmacovigilance concept extraction, "
                "we conducted an ablation study comparing four transformer encoder architectures (BERT-Base, BioBERT, Bio_ClinicalBERT, and ClinicalBERT) "
                "on the VAERS dataset across 5 independent random initialization seeds (Table 8). BioBERT achieved the highest validation F1 "
                "(0.8471 ± 0.0058, Precision = 0.8666, Recall = 0.8285), confirming its selection as the primary encoder architecture. "
                "All four models demonstrated close performance with minimal variance across seeds (SD < 0.009)."
            )
        elif txt.startswith("Table 8.") and "Pretrained Transformer Encoder" in txt:
            p.text = "Table 8. Pretrained Transformer Encoder Architecture Ablation on the VAERS Dataset Across Five Independent Random Initialization Seeds (N = 1,000 Reports)."

    # 4. Update Images (Figures 2-6)
    img_map = {
        "word/media/image2.png": manuscript_dir / "figure2.png",
        "word/media/image3.png": manuscript_dir / "figure3.png",
        "word/media/image4.png": manuscript_dir / "figure4.png",
        "word/media/image5.png": manuscript_dir / "figure5.png",
        "word/media/image6.png": manuscript_dir / "figure6.png",
    }

    try:
        doc.save(str(dst_docx_path))
        print(f"Saved directly to {dst_docx_path}")
        replace_media_images_in_memory(dst_docx_path, img_map)
    except PermissionError:
        alt_path = manuscript_dir / "LLM4AE_rev1_clean_updated2.docx"
        doc.save(str(alt_path))
        print(f"PermissionError on {dst_docx_path}. Saved directly to alternative path: {alt_path}")
        replace_media_images_in_memory(alt_path, img_map)


if __name__ == "__main__":
    main()
