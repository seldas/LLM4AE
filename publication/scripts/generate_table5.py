#!/usr/bin/env python3
"""Generate publication-ready Table 5: Leave-One-Drug-Event-Pair-Out Cross-Validation Performance on FAERS.

Evaluates Supervised BioBERT generalization across 4 Drug-AE Case Series (N = 829 Reports Total)
on all 17 clinical concept categories under a 4-Fold Leave-One-Drug-Event-Pair-Out protocol across 5 random seeds.

Metrics:
- 4 Individual Drug-AE Case Series: Mean ± SD across 5 seeds
- Overall Total: Aggregated Micro-Average across the 4 cohorts for each seed (Pooled N = 829 reports),
  reported as Mean ± SD across the 5 seeds.

All numbers are computed directly from raw evaluator outputs (publication/results/bert_runs_FAERS_LOO/raw.xlsx).

Outputs:
- publication/results/tables/table5_leave_one_out_faers.md
- publication/manuscripts/Tables/table5.md
- publication/manuscripts/table5.md
- publication/results/tables/table5_leave_one_out_faers.xlsx
- publication/results/tables/table5_data.json
- Updates Table 5 in publication/manuscripts/LLM4AE_rev1.docx
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
from pathlib import Path
from typing import Any, Dict, List, Tuple

import pandas as pd
import docx


PROJECT_ROOT = Path(__file__).resolve().parent.parent
BERT_RAW_PATH_DEFAULT = PROJECT_ROOT / "results" / "bert_runs_FAERS_LOO" / "raw.xlsx"
TABLES_DIR_DEFAULT = PROJECT_ROOT / "results" / "tables"
MANUSCRIPT_DIR_DEFAULT = PROJECT_ROOT / "manuscripts"
DOCX_PATH_DEFAULT = MANUSCRIPT_DIR_DEFAULT / "LLM4AE_rev1.docx"

FOLD_CONFIG = [
    ("Azacitidine – QT Prolongation", "Azacitidine-QT", "N = 200 reports", 200),
    ("Baricitinib – Hypersensitivity", "Baricitinib-Hypersensitivity", "N = 200 reports", 200),
    ("Tramadol – Hypoglycemia", "Tramadol-Hypoglycemia", "N = 229 reports", 229),
    ("Erenumab – Stroke", "Erenumab-Stroke", "N = 200 reports", 200),
]


def calculate_two_tier_metrics(df: pd.DataFrame) -> Dict[str, float]:
    """Compute Strict Exact-Match F1 and Adapted ADE-Eval F1 from raw aligned rows."""
    if df.empty:
        return {"strict_F1": 0.0, "ade_F1": 0.0, "strict_P": 0.0, "strict_R": 0.0, "ade_P": 0.0, "ade_R": 0.0}

    counts = df["match_type"].value_counts().to_dict()
    M = int(counts.get("M", 0))
    class_confusion_mask = (df["match_type"] == "C") & (df["error_subtype"] == "class_confusion") if "error_subtype" in df.columns else pd.Series(False, index=df.index)
    C_class = int(class_confusion_mask.sum())
    C_total = int(counts.get("C", 0))
    C_boundary = C_total - C_class
    S_non_overlap = int(counts.get("S", 0))
    N = int(counts.get("N", 0))

    # Scheme 3: Strict Exact-Match
    p3_den = M + C_total + S_non_overlap
    r3_den = M + C_total + N
    p3 = M / p3_den if p3_den > 0 else 0.0
    r3 = M / r3_den if r3_den > 0 else 0.0
    f3 = 2 * p3 * r3 / (p3 + r3) if (p3 + r3) > 0 else 0.0

    # Scheme 2: Adapted ADE-Eval
    m2 = M + 0.5 * C_total
    p2_den = M + C_total + 0.25 * S_non_overlap
    r2_den = M + C_total + N
    p2 = m2 / p2_den if p2_den > 0 else 0.0
    r2 = m2 / r2_den if r2_den > 0 else 0.0
    f2 = 2 * p2 * r2 / (p2 + r2) if (p2 + r2) > 0 else 0.0

    return {
        "strict_P": round(p3, 4),
        "strict_R": round(r3, 4),
        "strict_F1": round(f3, 4),
        "ade_P": round(p2, 4),
        "ade_R": round(r2, 4),
        "ade_F1": round(f2, 4),
    }


def compute_table5_data(raw_excel_path: Path) -> Tuple[List[dict], dict]:
    """Load raw LOO evaluation records and compute fold-level and micro-average aggregated metrics."""
    print(f"Loading raw BioBERT LOO records from: {raw_excel_path}")
    df_raw = pd.read_excel(raw_excel_path, sheet_name="Raw_Results")

    seeds = sorted(int(s) for s in df_raw["seed"].dropna().unique())
    print(f"Discovered {len(seeds)} random seeds: {seeds}")

    table_rows = []
    audit_data: Dict[str, Any] = {
        "seeds": seeds,
        "folds": {},
        "micro_average_pooled": {},
    }

    # 1. Individual 4 Folds
    for display_name, fold_key, cohort_desc, n_docs in FOLD_CONFIG:
        strict_list = []
        ade_list = []
        seed_metrics = {}

        for seed in seeds:
            fold_sub = df_raw[(df_raw["fold_name"] == fold_key) & (df_raw["seed"] == seed)]
            met = calculate_two_tier_metrics(fold_sub)
            strict_list.append(met["strict_F1"])
            ade_list.append(met["ade_F1"])
            seed_metrics[str(seed)] = met

        strict_s = pd.Series(strict_list)
        ade_s = pd.Series(ade_list)

        s_mean, s_std = strict_s.mean(), strict_s.std()
        a_mean, a_std = ade_s.mean(), ade_s.std()

        table_rows.append({
            "Drug–Event Case Series": display_name,
            "Validation Cohort Size": cohort_desc,
            "Strict Exact F1": f"{s_mean:.4f} ± {s_std:.4f}",
            "Adapted ADE F1": f"{a_mean:.4f} ± {a_std:.4f}",
            "_strict_mean": s_mean,
            "_strict_std": s_std,
            "_ade_mean": a_mean,
            "_ade_std": a_std,
            "_is_total": False,
        })

        audit_data["folds"][fold_key] = {
            "display_name": display_name,
            "cohort_size": cohort_desc,
            "n_docs": n_docs,
            "strict_F1_mean": round(float(s_mean), 4),
            "strict_F1_std": round(float(s_std), 4),
            "ade_F1_mean": round(float(a_mean), 4),
            "ade_F1_std": round(float(a_std), 4),
            "per_seed_results": seed_metrics,
        }

    # 2. Aggregated Micro-Average across all 4 folds for each seed
    pooled_strict_list = []
    pooled_ade_list = []
    pooled_seed_metrics = {}

    for seed in seeds:
        seed_sub = df_raw[df_raw["seed"] == seed]
        met = calculate_two_tier_metrics(seed_sub)
        pooled_strict_list.append(met["strict_F1"])
        pooled_ade_list.append(met["ade_F1"])
        pooled_seed_metrics[str(seed)] = met

    p_strict_s = pd.Series(pooled_strict_list)
    p_ade_s = pd.Series(pooled_ade_list)

    p_s_mean, p_s_std = p_strict_s.mean(), p_strict_s.std()
    p_a_mean, p_a_std = p_ade_s.mean(), p_ade_s.std()

    table_rows.append({
        "Drug–Event Case Series": "Total (Micro-Average Aggregated)",
        "Validation Cohort Size": "N = 829 reports total",
        "Strict Exact F1": f"{p_s_mean:.4f} ± {p_s_std:.4f}",
        "Adapted ADE F1": f"{p_a_mean:.4f} ± {p_a_std:.4f}",
        "_strict_mean": p_s_mean,
        "_strict_std": p_s_std,
        "_ade_mean": p_a_mean,
        "_ade_std": p_a_std,
        "_is_total": True,
    })

    audit_data["micro_average_pooled"] = {
        "description": "Micro-average pooled across all 4 case series (829 FAERS reports total) for each seed",
        "strict_F1_mean": round(float(p_s_mean), 4),
        "strict_F1_std": round(float(p_s_std), 4),
        "ade_F1_mean": round(float(p_a_mean), 4),
        "ade_F1_std": round(float(p_a_std), 4),
        "per_seed_results": pooled_seed_metrics,
    }

    return table_rows, audit_data


def generate_markdown_content(table_rows: List[dict]) -> str:
    """Format Table 5 rows into clean GitHub-flavored markdown with footnotes."""
    lines = [
        "# Table 5: Leave-One-Drug-Event-Pair-Out Cross-Validation Performance Across Four FAERS Case Series (N = 829 Reports Total)",
        "",
        "Supervised BioBERT model generalization evaluated under a 4-fold Leave-One-Drug-Event-Pair-Out cross-validation protocol on all 17 clinical concept categories. For each case series, the model was trained on the remaining 3 case series and evaluated on the held-out target series across 5 independent random initialization seeds.",
        "",
        "| Drug–Event Case Series | Validation Cohort Size | Primary Tier: Strict Exact F1 | Secondary Tier: Adapted ADE F1 |",
        "| :--- | :---: | :---: | :---: |",
    ]

    for r in table_rows:
        is_tot = r["_is_total"]
        name = f"**{r['Drug–Event Case Series']}**" if is_tot else r["Drug–Event Case Series"]
        cohort = f"**{r['Validation Cohort Size']}**" if is_tot else r["Validation Cohort Size"]
        s_f1 = f"**{r['Strict Exact F1']}**" if is_tot else r["Strict Exact F1"]
        a_f1 = f"**{r['Adapted ADE F1']}**" if is_tot else r["Adapted ADE F1"]
        lines.append(f"| {name} | {cohort} | {s_f1} | {a_f1} |")

    tot_row = table_rows[-1]
    lines.extend([
        "",
        "---",
        "",
        "### Footnotes & Methodological Notes:",
        "1. **Validation Design:** In each fold, all cases of a specific drug-event pair were completely held out from training to simulate real-world pharmacovigilance surveillance for emerging adverse drug reactions.",
        "2. **Evaluation Metrics:** Evaluated across the full 17 clinical concept categories. Mean $\\pm$ SD reflects variance across 5 independent training runs per case series ($N = 20$ total model runs).",
        f"3. **Micro-Average Aggregation:** The overall Total row reflects micro-average aggregation pooled across all 829 reports for each random seed ({tot_row['Strict Exact F1']} Strict Exact F1, {tot_row['Adapted ADE F1']} Adapted ADE F1).",
        "",
    ])
    return "\n".join(lines)


def update_docx_manuscript(docx_path: Path, table_rows: List[dict]) -> None:
    """Update Table 5 and associated narrative text in LLM4AE_rev1.docx."""
    if not docx_path.exists():
        print(f"Warning: docx manuscript not found at {docx_path}")
        return

    doc = docx.Document(str(docx_path))
    
    # Locate Table 5 (Table 4 in doc.tables index)
    target_table = None
    for t in doc.tables:
        header_text = " ".join(c.text for c in t.rows[0].cells).lower()
        if "drug" in header_text and "case series" in header_text and "validation cohort" in header_text:
            target_table = t
            break

    if target_table is not None:
        print(f"Found Table 5 in {docx_path.name}; updating table cells...")
        # Map row data
        for row_idx, r_data in enumerate(table_rows, start=1):
            if row_idx < len(target_table.rows):
                row_cells = target_table.rows[row_idx].cells
                row_cells[0].text = r_data["Drug–Event Case Series"].replace("Total (Micro-Average Aggregated)", "Total (Micro-Average)")
                row_cells[1].text = r_data["Validation Cohort Size"].replace(" reports", "").replace(" total", "").replace("N = ", "")
                row_cells[2].text = r_data["Strict Exact F1"]
                row_cells[3].text = r_data["Adapted ADE F1"]

    # Update narrative text paragraph referring to Table 5
    tot_row = table_rows[-1]
    p_updated = False
    for p in doc.paragraphs:
        if "Table 5" in p.text and "azacitidine" in p.text.lower() and "erenumab" in p.text.lower():
            p.text = (
                "We applied a repeated Leave-One-Drug-Event-Pair-Out cross-validation protocol on the FAERS corpus to evaluate "
                "model generalization across distinct therapeutic contexts on all 17 clinical concept categories. For each of "
                "the 4 curated drug-event cohorts, the model was trained on the remaining 3 case series and evaluated on the held-out "
                f"series. As shown in Table 5, strict exact-match F1 was {table_rows[0]['_strict_mean']:.4f} ± {table_rows[0]['_strict_std']:.4f} for azacitidine–QT prolongation (N = 200), "
                f"{table_rows[1]['_strict_mean']:.4f} ± {table_rows[1]['_strict_std']:.4f} for baricitinib–hypersensitivity (N = 200), "
                f"{table_rows[2]['_strict_mean']:.4f} ± {table_rows[2]['_strict_std']:.4f} for tramadol–hypoglycemia (N = 229), and "
                f"{table_rows[3]['_strict_mean']:.4f} ± {table_rows[3]['_strict_std']:.4f} for erenumab–stroke (N = 200), yielding an aggregated "
                f"micro-average F1 of {tot_row['Strict Exact F1']} across the full corpus ({tot_row['Adapted ADE F1']} Adapted ADE F1). "
                "The between-series performance variance exceeded within-fold variation, demonstrating that narrative complexity "
                "and therapeutic vocabulary differences contribute more variation than stochastic model initialization."
            )
            p_updated = True
            break

    try:
        doc.save(str(docx_path))
        print(f"Successfully updated Table 5 and narrative in: {docx_path}")
    except PermissionError:
        alt_path = docx_path.with_name(f"{docx_path.stem}_table5_updated{docx_path.suffix}")
        doc.save(str(alt_path))
        print(
            f"Notice: '{docx_path.name}' is currently open in Microsoft Word (locked).\n"
            f"Saved updated version to: '{alt_path}'.\n"
            f"Please close the document in Word to allow direct overwrite."
        )
    except Exception as e:
        print(f"Warning: Could not save updated docx: {e}")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--raw-excel", type=Path, default=BERT_RAW_PATH_DEFAULT, help="Path to bert_runs_FAERS_LOO/raw.xlsx")
    parser.add_argument("--tables-dir", type=Path, default=TABLES_DIR_DEFAULT, help="Output directory for table artifacts")
    parser.add_argument("--manuscript-dir", type=Path, default=MANUSCRIPT_DIR_DEFAULT, help="Path to manuscripts directory")
    parser.add_argument("--docx-path", type=Path, default=DOCX_PATH_DEFAULT, help="Path to LLM4AE_rev1.docx")
    args = parser.parse_args()

    args.tables_dir.mkdir(parents=True, exist_ok=True)
    args.manuscript_dir.mkdir(parents=True, exist_ok=True)

    table_rows, audit_data = compute_table5_data(args.raw_excel)

    # 1. Print Console Summary
    print("\n=== Table 5: Leave-One-Drug-Event-Pair-Out Cross-Validation Performance ===")
    for r in table_rows:
        print(f"{r['Drug–Event Case Series']:35s} | {r['Validation Cohort Size']:22s} | Strict: {r['Strict Exact F1']:17s} | Adapted: {r['Adapted ADE F1']:17s}")

    # 2. Markdown Files
    md_content = generate_markdown_content(table_rows)
    md_paths = [
        args.tables_dir / "table5_leave_one_out_faers.md",
        args.manuscript_dir / "Tables" / "table5.md",
        args.manuscript_dir / "table5.md",
    ]
    for p in md_paths:
        p.parent.mkdir(parents=True, exist_ok=True)
        with p.open("w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"Saved markdown: {p}")

    # 3. Excel Workbook with ESM Cover Sheet
    esm_cover = pd.DataFrame([
        {"Metadata Field": "Article Title", "Value": "Benchmarking Fine-Tuned Encoders and Instruction-Tuned Large Language Models for Adverse Event Clinical Concept Extraction from Spontaneous Reporting Narratives"},
        {"Metadata Field": "Journal", "Value": "Drug Safety"},
        {"Metadata Field": "Table Identifier", "Value": "Table 5: Leave-One-Drug-Event-Pair-Out Cross-Validation Performance on FAERS"},
        {"Metadata Field": "Corpus", "Value": "FDA Adverse Event Reporting System (FAERS, N = 829 Reports across 4 Case Series)"},
        {"Metadata Field": "Evaluation Protocol", "Value": "4-Fold Leave-One-Drug-Event-Pair-Out across 5 Random Seeds (N = 20 Model Runs)"},
        {"Metadata Field": "Aggregation Method", "Value": "Micro-Average Aggregated across all 829 reports for total row"},
        {"Metadata Field": "Generated By", "Value": "publication/scripts/generate_table5.py (Computed dynamically from raw.xlsx)"}
    ])

    export_df = pd.DataFrame([
        {
            "Drug–Event Case Series": r["Drug–Event Case Series"],
            "Validation Cohort Size": r["Validation Cohort Size"],
            "Primary Tier: Strict Exact F1": r["Strict Exact F1"],
            "Secondary Tier: Adapted ADE F1": r["Adapted ADE F1"],
        }
        for r in table_rows
    ])

    out_excel = args.tables_dir / "table5_leave_one_out_faers.xlsx"
    with pd.ExcelWriter(out_excel, engine="openpyxl") as writer:
        esm_cover.to_excel(writer, sheet_name="ESM_Cover_Sheet", index=False)
        export_df.to_excel(writer, sheet_name="Table_5_FAERS_LOO", index=False)
    print(f"Saved excel: {out_excel}")

    # 4. JSON Audit Data
    out_json = args.tables_dir / "table5_data.json"
    with out_json.open("w", encoding="utf-8") as f:
        json.dump(audit_data, f, indent=2)
    print(f"Saved audit json: {out_json}")

    # 5. Update Word Manuscript Docx
    update_docx_manuscript(args.docx_path, table_rows)


if __name__ == "__main__":
    main()
