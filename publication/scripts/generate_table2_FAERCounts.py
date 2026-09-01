#!/usr/bin/env python3
"""Generate Table 2 (FAERS Annotation Counts by Category) for LLM4AE.

Calculates annotation counts for the 17 clinical concept categories across:
1. Human Ground Truth (SME1 annotations from dataset.db, with bSYM merged into Dx)
2. ETHER rule-based baseline (from dataset.db)
3. LLaMA-4 predictions (from results/llama4_runs_FAERS/llama4_raw.xlsx)
4. Claude 4.6 Sonnet predictions (from results/sonnet_runs_FAERS/sonnet_raw.xlsx)

Updates publication/manuscripts/LLM4AE_rev1.docx Table 2 and exports
summary markdown and excel files.
"""

from __future__ import annotations

import argparse
import json
import sqlite3
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import docx
import pandas as pd


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DB_PATH_DEFAULT = PROJECT_ROOT / "dataset.db"
DOCX_PATH_DEFAULT = PROJECT_ROOT / "manuscripts" / "LLM4AE_rev1.docx"
RESULTS_DIR_DEFAULT = PROJECT_ROOT / "results"
TABLES_DIR_DEFAULT = RESULTS_DIR_DEFAULT / "tables"

# 17 standardized clinical concept categories in canonical presentation order
CATEGORIES = [
    ("SDrug", "Suspect Drug"),
    ("CDrug", "Concomitant Drug"),
    ("ODrug", "Other Drug"),
    ("Dose", "Dosage"),
    ("Indication", "Drug Indication"),
    ("Treatment", "Drug used for Treatment"),
    ("AE", "Adverse Event"),
    ("mAE", "AE Manifestation/Sequela"),
    ("Dx", "Diagnostic Test / Baseline Symptom"),
    ("Lab", "Laboratory Finding"),
    ("Status", "Patient Clinical Status"),
    ("R/O", "Rule-Out Diagnosis"),
    ("CoD", "Cause of Death"),
    ("MHx", "Medical History"),
    ("FHx", "Family History"),
    ("Age", "Patient Age"),
    ("Sex", "Patient Sex"),
]

# Raw human (SME1) label mapping -> standard category (with bSYM merged into Dx)
HUMAN_LABEL_MAP = {
    "sdrug": "SDrug",
    "cdrug": "CDrug",
    "drug": "ODrug",
    "odrug": "ODrug",
    "dose": "Dose",
    "indication": "Indication",
    "treatment": "Treatment",
    "ae": "AE",
    "mae": "mAE",
    "diagnostic": "Dx",
    "dx": "Dx",
    "bsym": "Dx",  # Normalized bSYM -> Dx
    "baseline symptom": "Dx",  # Normalized bSYM -> Dx
    "lab": "Lab",
    "status": "Status",
    "r/o": "R/O",
    "ro": "R/O",
    "cause of death": "CoD",
    "cod": "CoD",
    "medical history": "MHx",
    "mhx": "MHx",
    "family history": "FHx",
    "fhx": "FHx",
    "age": "Age",
    "sex": "Sex",
}

# Raw LLM label mapping -> standard category
LLM_LABEL_MAP = {
    "sdrug": "SDrug",
    "cdrug": "CDrug",
    "odrug": "ODrug",
    "drug": "ODrug",
    "dose": "Dose",
    "indication": "Indication",
    "treatment": "Treatment",
    "ae": "AE",
    "mae": "mAE",
    "diagnostic": "Dx",
    "dx": "Dx",
    "bsym": "Dx",
    "lab": "Lab",
    "status": "Status",
    "ro": "R/O",
    "r/o": "R/O",
    "cod": "CoD",
    "mhx": "MHx",
    "fhx": "FHx",
    "age": "Age",
    "sex": "Sex",
}


def load_human_counts(db_path: Path) -> Dict[str, int]:
    """Load FAERS SME1 ground truth counts from SQLite database."""
    counts: Dict[str, int] = Counter()
    with sqlite3.connect(db_path) as conn:
        rows = conn.execute("""
            SELECT a.label
            FROM annotations a
            JOIN documents d ON a.doc_id = d.doc_id
            WHERE d.dataset = 'FAERS' AND a.note = 'SME1'
        """).fetchall()

    for (raw_label,) in rows:
        key = str(raw_label).strip().lower()
        mapped = HUMAN_LABEL_MAP.get(key)
        if mapped:
            counts[mapped] += 1
    return dict(counts)


def load_ether_counts(db_path: Path) -> Dict[str, str]:
    """Load FAERS ETHER baseline counts from SQLite database.
    
    ETHER uses broader umbrella categories:
    - DRUG / VACCINE (17,164) across SDrug, CDrug, ODrug
    - SYMPTOM (25,031) across AE, mAE
    - DIAGNOSIS / SECOND_LEVEL_DIAGNOSIS (4,344) for Dx
    - RULE_OUT (33) for R/O
    - CAUSE_OF_DEATH (125) for CoD
    - MEDICAL_HISTORY (755) for MHx
    - FAMILY_HISTORY (22) for FHx
    """
    raw_counts: Dict[str, int] = Counter()
    with sqlite3.connect(db_path) as conn:
        rows = conn.execute("""
            SELECT a.label
            FROM annotations a
            JOIN documents d ON a.doc_id = d.doc_id
            WHERE d.dataset = 'FAERS' AND a.note = 'ETHER'
        """).fetchall()

    for (raw_label,) in rows:
        raw_counts[str(raw_label).strip()] += 1

    drug_total = raw_counts.get("DRUG", 0) + raw_counts.get("VACCINE", 0)
    symptom_total = raw_counts.get("SYMPTOM", 0)
    dx_total = raw_counts.get("DIAGNOSIS", 0) + raw_counts.get("SECOND_LEVEL_DIAGNOSIS", 0)
    ro_total = raw_counts.get("RULE_OUT", 0)
    cod_total = raw_counts.get("CAUSE_OF_DEATH", 0)
    mhx_total = raw_counts.get("MEDICAL_HISTORY", 0)
    fhx_total = raw_counts.get("FAMILY_HISTORY", 0)

    ether_map: Dict[str, str] = {
        "SDrug": f"{drug_total:,}*",
        "CDrug": f"{drug_total:,}*",
        "ODrug": f"{drug_total:,}*",
        "Dose": "n/a",
        "Indication": "n/a",
        "Treatment": "n/a",
        "AE": f"{symptom_total:,}*",
        "mAE": f"{symptom_total:,}*",
        "Dx": f"{dx_total:,}",
        "Lab": "n/a",
        "Status": "n/a",
        "R/O": f"{ro_total:,}",
        "CoD": f"{cod_total:,}",
        "MHx": f"{mhx_total:,}",
        "FHx": f"{fhx_total:,}",
        "Age": "n/a",
        "Sex": "n/a",
    }
    return ether_map


def load_llm_counts_from_excel_or_jsonl(raw_xlsx_path: Path, jsonl_path: Path) -> Dict[str, int]:
    """Load model prediction counts from raw.xlsx or predictions.jsonl."""
    counts: Dict[str, int] = Counter()
    if raw_xlsx_path.exists():
        df = pd.read_excel(raw_xlsx_path)
        pred_series = df.loc[df["label_pred"].notna(), "label_pred"]
        for raw_label in pred_series:
            key = str(raw_label).strip().lower()
            mapped = LLM_LABEL_MAP.get(key)
            if mapped:
                counts[mapped] += 1
    elif jsonl_path.exists():
        with jsonl_path.open(encoding="utf-8") as f:
            for line in f:
                if not line.strip():
                    continue
                doc = json.loads(line)
                for ent in doc.get("predicted_entities", []):
                    key = str(ent.get("label", "")).strip().lower()
                    mapped = LLM_LABEL_MAP.get(key)
                    if mapped:
                        counts[mapped] += 1
    else:
        raise FileNotFoundError(f"Neither {raw_xlsx_path} nor {jsonl_path} exists.")
    return dict(counts)


def build_table2_dataframe(
    human_counts: Dict[str, int],
    ether_counts: Dict[str, str],
    llama_counts: Dict[str, int],
    sonnet_counts: Dict[str, int],
) -> pd.DataFrame:
    """Construct structured DataFrame for Table 2."""
    rows = []
    for cat_key, _ in CATEGORIES:
        h_val = human_counts.get(cat_key, 0)
        e_val = ether_counts.get(cat_key, "n/a")
        l_val = llama_counts.get(cat_key, 0)
        s_val = sonnet_counts.get(cat_key, 0)
        rows.append({
            "Category": cat_key,
            "Human": f"{h_val:,}",
            "ETHER": e_val,
            "LLM (LLaMA-4)": f"{l_val:,}",
            "LLM (Sonnet 4.6)": f"{s_val:,}",
        })
    return pd.DataFrame(rows)


def update_manuscript_docx(docx_path: Path, table_df: pd.DataFrame) -> None:
    """Update Table 2 in LLM4AE_rev1.docx while preserving document formatting."""
    doc = docx.Document(docx_path)
    target_table = None

    for table in doc.tables:
        header_text = [cell.text.strip() for cell in table.rows[0].cells]
        if "Category" in header_text and "Human" in header_text and "ETHER" in header_text:
            target_table = table
            break

    if target_table is None:
        raise RuntimeError("Table 2 (FAERS annotations) not found in Word document.")

    category_to_row = {row["Category"]: row for _, row in table_df.iterrows()}

    for row_idx in range(1, len(target_table.rows)):
        row = target_table.rows[row_idx]
        cat_cell_text = row.cells[0].text.strip()
        # Find matching category key
        matched_cat = None
        for cat_key, _ in CATEGORIES:
            if cat_cell_text.lower() == cat_key.lower() or cat_cell_text.replace(" ", "").lower() == cat_key.lower():
                matched_cat = cat_key
                break

        if matched_cat and matched_cat in category_to_row:
            data = category_to_row[matched_cat]
            values = [
                data["Category"],
                data["Human"],
                data["ETHER"],
                data["LLM (LLaMA-4)"],
                data["LLM (Sonnet 4.6)"],
            ]
            for col_idx, val in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = str(val)
                # Ensure font formatting
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = docx.shared.Pt(8)

    # Clean footnote in paragraphs if any contains unencoded characters
    for p in doc.paragraphs:
        if "ETHER does not distinguish drug roles" in p.text:
            p.text = "*ETHER does not distinguish drug roles or AE types – all drug and AE mentions are output under a single DRUG tag (17,164) and AE tag (25,031)."
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = docx.shared.Pt(9)

    doc.save(docx_path)
    print(f"Successfully updated Table 2 in {docx_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--db-path", type=Path, default=DB_PATH_DEFAULT, help="Path to dataset.db")
    parser.add_argument("--docx-path", type=Path, default=DOCX_PATH_DEFAULT, help="Path to LLM4AE_rev1.docx")
    parser.add_argument("--results-dir", type=Path, default=RESULTS_DIR_DEFAULT, help="Path to results directory")
    parser.add_argument("--update-docx", action="store_true", default=True, help="Update LLM4AE_rev1.docx Table 2")
    parser.add_argument("--no-docx", dest="update_docx", action="store_false", help="Skip docx update")
    args = parser.parse_args()

    print(f"Loading data from SQLite database: {args.db_path}")
    human_counts = load_human_counts(args.db_path)
    ether_counts = load_ether_counts(args.db_path)

    llama_dir = args.results_dir / "llama4_runs_FAERS"
    sonnet_dir = args.results_dir / "sonnet_runs_FAERS"

    print(f"Loading LLaMA-4 predictions from: {llama_dir}")
    llama_counts = load_llm_counts_from_excel_or_jsonl(
        llama_dir / "llama4_raw.xlsx", llama_dir / "predictions.jsonl"
    )

    print(f"Loading Sonnet 4.6 predictions from: {sonnet_dir}")
    sonnet_counts = load_llm_counts_from_excel_or_jsonl(
        sonnet_dir / "sonnet_raw.xlsx", sonnet_dir / "predictions.jsonl"
    )

    table_df = build_table2_dataframe(human_counts, ether_counts, llama_counts, sonnet_counts)

    print("\n" + "=" * 70)
    print("Table 2: FAERS Annotations (bSYM merged into Dx)")
    print("=" * 70)
    print(table_df.to_string(index=False))
    print("=" * 70)

    # Save to Excel & Markdown in tables directory
    tables_dir = args.results_dir / "tables"
    tables_dir.mkdir(parents=True, exist_ok=True)
    out_xlsx = tables_dir / "table2_FAERS_annotations.xlsx"
    out_md = tables_dir / "table2_FAERS_annotations.md"

    table_df.to_excel(out_xlsx, index=False)
    
    # Generate markdown table manually without tabulate dependency
    headers = list(table_df.columns)
    header_line = "| " + " | ".join(headers) + " |"
    sep_line = "| " + " | ".join([":---" if i == 0 else ":---:" for i in range(len(headers))]) + " |"
    data_lines = ["| " + " | ".join(str(val) for val in row) + " |" for row in table_df.values]
    table_md_str = "\n".join([header_line, sep_line] + data_lines)

    md_content = f"# Table 2: FAERS Annotations Breakdown\n\n{table_md_str}\n\n" \
                 f"*ETHER does not distinguish drug roles or AE types – all drug and AE mentions " \
                 f"are output under a single DRUG tag (17,164) and AE tag (25,031).\n"
    out_md.write_text(md_content, encoding="utf-8")
    print(f"Saved table outputs to:\n  {out_xlsx}\n  {out_md}")

    if args.update_docx and args.docx_path.exists():
        update_manuscript_docx(args.docx_path, table_df)


if __name__ == "__main__":
    main()

