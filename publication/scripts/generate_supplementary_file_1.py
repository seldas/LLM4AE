#!/usr/bin/env python3
"""
generate_supplementary_file_1.py

Generates Supplementary File 1 (Prompts) in both Word (.docx) and Markdown (.md) formats.
Includes all 4 primary LLM prompt templates without the embedded annotation guidelines table:
1. FAERS In-Text XML Tagging Prompt (P2_TAG)
2. FAERS Structured JSON Schema Prompt (P1_JSON)
3. VAERS In-Text XML Tagging Prompt (P2_TAG_VAERS)
4. VAERS Structured JSON Schema Prompt (P1_JSON_VAERS)
"""

from __future__ import annotations

from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


# ==============================================================================
# PROMPT DEFINITIONS (Excluding the extensive Annotation Guide table)
# ==============================================================================

PROMPT_FAERS_P2_TAG = r'''You are an expert medical annotator analyzing a FAERS (FDA Adverse Event Reporting System) case report narrative.

Your task is to identify clinical entities according to the annotation schema and insert XML-style annotation tags directly into the original narrative.

### Allowed Tags

Use ONLY these tags:

<SDRUG>...</SDRUG>
<CDRUG>...</CDRUG>
<ODRUG>...</ODRUG>
<DOSE>...</DOSE>
<IND>...</IND>
<TREATMENT>...</TREATMENT>
<AE>...</AE>
<MAE>...</MAE>
<DX>...</DX>
<LAB>...</LAB>
<STATUS>...</STATUS>
<RO>...</RO>
<COD>...</COD>
<MHX>...</MHX>
<FHX>...</FHX>
<AGE>...</AGE>
<SEX>...</SEX>

Do NOT create any other tag.

### In-Text Annotation Rules

1. Insert tags around the exact entity span in the original narrative.

2. Do NOT alter the original narrative in any way other than inserting annotation tags.

3. Preserve exactly:
   - wording
   - spelling
   - capitalization
   - punctuation
   - numbers
   - whitespace
   - paragraph structure

4. Every opening tag must have the corresponding closing tag.

5. Tags must NOT overlap or nest.

6. Annotate only the smallest complete clinically meaningful span.

7. Contextual or trigger phrases should normally remain outside the tag.

### Examples

Original:
The patient was treated with prednisone for rash.

Correct:
The patient was treated with <TREATMENT>prednisone</TREATMENT> for <IND>rash</IND>.

Incorrect:
The patient was <TREATMENT>treated with prednisone</TREATMENT> for rash.

Original:
Concomitant medications included atenolol 25 mg twice daily.

Correct:
Concomitant medications included <CDRUG>atenolol</CDRUG> <DOSE>25 mg twice daily</DOSE>.

Original:
Her medical history included hypertension.

Correct:
Her medical history included <MHX>hypertension</MHX>.

Original:
CT demonstrated no acute intracranial abnormality.

Correct:
<DX>CT</DX> demonstrated no acute intracranial abnormality.

### Narrative

{text}

### CRITICAL OUTPUT REQUIREMENTS

1. Return ONLY the fully annotated narrative.
2. Do NOT add an introductory sentence such as "The annotated text is shown as below:".
3. Do NOT use Markdown code fences.
4. Do NOT provide explanations, comments, summaries, or lists.
5. Apart from the inserted annotation tags, every character of the original narrative must remain unchanged.'''


PROMPT_FAERS_P1_JSON = r'''You are an expert medical annotator analyzing a FAERS (FDA Adverse Event Reporting System) case report narrative.

Your task is to identify clinical entities in the narrative according to the annotation schema and return the annotations as structured JSON.

### JSON Output Schema

Return exactly one JSON object containing all 17 keys below:

{
  "sdrug": [],
  "cdrug": [],
  "odrug": [],
  "dose": [],
  "ind": [],
  "treatment": [],
  "ae": [],
  "mae": [],
  "dx": [],
  "lab": [],
  "status": [],
  "ro": [],
  "cod": [],
  "mhx": [],
  "fhx": [],
  "age": [],
  "sex": []
}

Each detected entity must be represented as:

{
  "text": "exact substring from narrative",
  "start": 0,
  "end": 0
}

### Rules for "text", "start", and "end"

- "text" MUST be copied verbatim from the narrative.
- "start" MUST be the 0-based character offset of the first character of "text" in the supplied narrative.
- "end" MUST be the 0-based exclusive character offset immediately after the last character of "text".
- The intended relationship is: narrative[start:end] == text.
- Count every character exactly as it appears in the supplied narrative, including spaces, punctuation, and newline characters.
- Do not normalize, rewrite, expand, abbreviate, or correct the text.
- Do not include unnecessary contextual words around the entity.
- If the same entity text occurs multiple times, use the offsets of the specific occurrence being annotated.
- Each explicit occurrence must be represented separately.
- Within each category, order entities by ascending "start", then ascending "end".

### Completeness and Ordering Rules

- Include every supported entity occurrence found in the narrative.
- Repeated occurrences must be returned as separate objects.
- Do not collapse repeated mentions into a single object.
- If a category has no entities, return an empty list.
- Return all 17 keys, even when their values are empty lists.
- Do not return duplicate objects for the same occurrence.
- Within each category list, order annotations by ascending "start", then ascending "end".

### Narrative

{text}

### CRITICAL OUTPUT REQUIREMENTS

1. Return ONLY valid JSON.
2. Do NOT use Markdown code fences.
3. Do NOT include ```json or ```.
4. Do NOT include explanations, headings, comments, or conversational text.
5. The first character of the response must be "{".
6. The final character of the response must be "}".'''


PROMPT_VAERS_P2_TAG = r'''You are an expert medical annotator analyzing a VAERS (Vaccine Adverse Event Reporting System) case report narrative.

Your task is to identify clinical and contextual entities according to the annotation schema and insert XML-style annotation tags directly into the original narrative.

### Allowed Tags

Use ONLY these tags:

<SYM>...</SYM>
<SDX>...</SDX>
<PDX>...</PDX>
<DX>...</DX>
<VAX>...</VAX>
<MHX>...</MHX>
<FHX>...</FHX>
<LAB>...</LAB>
<TEMPO>...</TEMPO>
<DOSE>...</DOSE>
<STATUS>...</STATUS>
<TX>...</TX>
<AGE>...</AGE>
<SEX>...</SEX>

Do NOT create any other tag.

### In-Text Annotation Rules

1. Insert tags around the exact entity span in the original narrative.

2. Do NOT alter the original narrative in any way other than inserting annotation tags.

3. Preserve exactly:
   - wording
   - spelling
   - capitalization
   - punctuation
   - numbers
   - whitespace
   - paragraph structure

4. Every opening tag must have the corresponding closing tag.

5. Tags must NOT overlap or nest.

6. Annotate only the smallest complete clinically meaningful span.

7. Contextual or trigger phrases should normally remain outside the tag.

### Examples

Original:
A 45-year-old female received the second dose of Pfizer COVID-19 vaccine and developed fever and headache the next day.

Correct:
A <AGE>45-year-old</AGE> <SEX>female</SEX> received the <DOSE>second dose</DOSE> of <VAX>Pfizer COVID-19 vaccine</VAX> and developed <SYM>fever</SYM> and <SYM>headache</SYM> <TEMPO>the next day</TEMPO>.

Original:
She was diagnosed with myocarditis and treated with ibuprofen.

Correct:
She was diagnosed with <SDX>myocarditis</SDX> and treated with <TX>ibuprofen</TX>.

Original:
The emergency physician was concerned for possible myocarditis.

Correct:
The emergency physician was concerned for possible <PDX>myocarditis</PDX>.

Original:
Past medical history included asthma.

Correct:
Past medical history included <MHX>asthma</MHX>.

Original:
Temperature was 39.1 C and heart rate was 112 bpm.

Correct:
<LAB>Temperature was 39.1 C</LAB> and <LAB>heart rate was 112 bpm</LAB>.

Original:
Symptoms resolved after two days and the patient was discharged home.

Correct:
Symptoms <STATUS>resolved</STATUS> <TEMPO>after two days</TEMPO> and the patient was <STATUS>discharged home</STATUS>.

### Narrative

{text}

### CRITICAL OUTPUT REQUIREMENTS

1. Return ONLY the fully annotated narrative.
2. Do NOT add an introductory sentence such as "The annotated text is shown as below:".
3. Do NOT use Markdown code fences.
4. Do NOT provide explanations, comments, summaries, or lists.
5. Apart from the inserted annotation tags, every character of the original narrative must remain unchanged.'''


PROMPT_VAERS_P1_JSON = r'''You are an expert medical annotator analyzing a VAERS (Vaccine Adverse Event Reporting System) case report narrative.

Your task is to identify clinical and contextual entities in the narrative according to the annotation schema and return the annotations as structured JSON.

### JSON Output Schema

Return exactly one JSON object containing all 14 keys below:

{
  "sym": [],
  "sdx": [],
  "pdx": [],
  "dx": [],
  "vax": [],
  "mhx": [],
  "fhx": [],
  "lab": [],
  "tempo": [],
  "dose": [],
  "status": [],
  "tx": [],
  "age": [],
  "sex": []
}

Each detected entity must be represented as:

{
  "text": "exact substring from narrative",
  "start": 0,
  "end": 0
}

### Rules for "text", "start", and "end"

- "text" MUST be copied verbatim from the narrative.
- "start" MUST be the 0-based character offset of the first character of "text" in the supplied narrative.
- "end" MUST be the 0-based exclusive character offset immediately after the last character of "text".
- The intended relationship is: narrative[start:end] == text.
- Count every character exactly as it appears in the supplied narrative, including spaces, punctuation, and newline characters.
- Do not normalize, rewrite, expand, abbreviate, or correct the text.
- Do not include unnecessary contextual words around the entity.
- If the same entity text occurs multiple times, use the offsets of the specific occurrence being annotated.
- Each explicit occurrence must be represented separately.
- Within each category, order entities by ascending "start", then ascending "end".

### Completeness and Ordering Rules

- Include every supported entity occurrence found in the narrative.
- Repeated occurrences must be returned as separate objects.
- Do not collapse repeated mentions into a single object.
- If a category has no entities, return an empty list.
- Return all 14 keys, even when their values are empty lists.
- Do not return duplicate objects for the same occurrence.
- Within each category list, order annotations by ascending "start", then ascending "end".

### Narrative

{text}

### CRITICAL OUTPUT REQUIREMENTS

1. Return ONLY valid JSON.
2. Do NOT use Markdown code fences.
3. Do NOT include ```json or ```.
4. Do NOT include explanations, headings, comments, or conversational text.
5. The first character of the response must be "{".
6. The final character of the response must be "}".'''


def add_code_block(doc, text: str):
    """Adds a bordered monospace code box to the Word document."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
            f'<w:left w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
            f'<w:right w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
            f'<w:insideH w:val="none"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
    tcPr.append(shd)

    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="120" w:type="dxa"/>'
        f'<w:bottom w:w="120" w:type="dxa"/>'
        f'<w:left w:w="140" w:type="dxa"/>'
        f'<w:right w:w="140" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

    cell.text = ""
    lines = text.strip().split("\n")
    for l_idx, line in enumerate(lines):
        p = cell.paragraphs[0] if l_idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = Pt(10.5)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def generate_docx(out_path: Path):
    doc = docx.Document()

    # Set Margins
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Supplementary File 1: Large Language Model Prompt Templates")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run(
        "Accompanying the manuscript: Benchmarking Large Language Models and Fine-Tuned Encoders for Clinical Concept Extraction "
        "from Pharmacovigilance and Vaccine Adverse Event Narratives"
    )
    run_sub.italic = True
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10.5)
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(14)
    r_intro = p_intro.add_run(
        "Note: This document contains the full system and task prompt instructions, allowed XML tags / JSON schema specifications, "
        "formatting rules, and few-shot examples utilized for prompting Large Language Models (LLMs) across the FAERS and VAERS benchmark evaluations. "
        "The extensive category definitions and annotation guidelines (which are programmatically injected at runtime) are provided in Supplementary File 2."
    )
    r_intro.font.name = "Arial"
    r_intro.font.size = Pt(9.5)

    # Section 1: FAERS
    h1 = doc.add_heading("1. FAERS Prompt Templates (17 Clinical Concept Categories)", level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    # Section 1.1: P2_TAG
    h2 = doc.add_heading("1.1 In-Text XML Tagging Prompt (P2_TAG)", level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("The prompt below specifies the inline XML tagging task (P2_TAG) for the 17 FAERS clinical concept categories:").paragraph_format.space_after = Pt(6)
    add_code_block(doc, PROMPT_FAERS_P2_TAG)

    # Section 1.2: P1_JSON
    h2 = doc.add_heading("1.2 Structured JSON Schema Prompt (P1_JSON)", level=2)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("The prompt below specifies the structured JSON schema task (P1_JSON) for the 17 FAERS clinical concept categories:").paragraph_format.space_after = Pt(6)
    add_code_block(doc, PROMPT_FAERS_P1_JSON)

    # Section 2: VAERS
    h1 = doc.add_heading("2. VAERS Prompt Templates (14 Clinical Concept Categories)", level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)

    # Section 2.1: P2_TAG_VAERS
    h2 = doc.add_heading("2.1 In-Text XML Tagging Prompt (P2_TAG_VAERS)", level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("The prompt below specifies the inline XML tagging task (P2_TAG_VAERS) for the 14 VAERS clinical concept categories:").paragraph_format.space_after = Pt(6)
    add_code_block(doc, PROMPT_VAERS_P2_TAG)

    # Section 2.2: P1_JSON_VAERS
    h2 = doc.add_heading("2.2 Structured JSON Schema Prompt (P1_JSON_VAERS)", level=2)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("The prompt below specifies the structured JSON schema task (P1_JSON_VAERS) for the 14 VAERS clinical concept categories:").paragraph_format.space_after = Pt(6)
    add_code_block(doc, PROMPT_VAERS_P1_JSON)

    doc.save(str(out_path))
    print(f"Saved docx to {out_path}")


def generate_markdown(out_path: Path):
    md_content = f"""# Supplementary File 1: Large Language Model Prompt Templates

**Accompanying the manuscript:**  
*Benchmarking Large Language Models and Fine-Tuned Encoders for Clinical Concept Extraction from Pharmacovigilance and Vaccine Adverse Event Narratives*

> **Note:** This document contains the full system and task prompt instructions, allowed XML tags / JSON schema specifications, formatting rules, and few-shot examples utilized for prompting Large Language Models (LLMs) across the FAERS and VAERS benchmark evaluations. The extensive category definitions and annotation guidelines (which are programmatically injected at runtime) are provided in Supplementary File 2.

---

## 1. FAERS Prompt Templates (17 Clinical Concept Categories)

### 1.1 In-Text XML Tagging Prompt (`P2_TAG`)

```text
{PROMPT_FAERS_P2_TAG}
```

---

### 1.2 Structured JSON Schema Prompt (`P1_JSON`)

```text
{PROMPT_FAERS_P1_JSON}
```

---

## 2. VAERS Prompt Templates (14 Clinical Concept Categories)

### 2.1 In-Text XML Tagging Prompt (`P2_TAG_VAERS`)

```text
{PROMPT_VAERS_P2_TAG}
```

---

### 2.2 Structured JSON Schema Prompt (`P1_JSON_VAERS`)

```text
{PROMPT_VAERS_P1_JSON}
```
"""
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"Saved markdown to {out_path}")


def main():
    repo_root = Path(__file__).resolve().parent.parent.parent
    manuscript_dir = repo_root / "publication" / "manuscripts"
    supp_dir = repo_root / "publication" / "supplementary"
    supp_dir.mkdir(parents=True, exist_ok=True)

    docx_out1 = manuscript_dir / "Supplementary_File_1_Prompts.docx"
    docx_out2 = supp_dir / "Supplementary_File_1_Prompts.docx"
    md_out1 = manuscript_dir / "Supplementary_File_1_Prompts.md"
    md_out2 = supp_dir / "Supplementary_File_1_Prompts.md"

    generate_docx(docx_out1)
    generate_docx(docx_out2)
    generate_markdown(md_out1)
    generate_markdown(md_out2)


if __name__ == "__main__":
    main()
